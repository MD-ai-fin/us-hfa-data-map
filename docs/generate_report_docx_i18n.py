#!/usr/bin/env python3
"""Build Word (.docx) source files for the ES/PL analysis reports.

The EN and ZH FY2025 analysis reports originate from hand-authored .docx
files (see docs/export_reports.py). The ES/PL versions were produced later
by translating the rendered HTML directly (docs/reports/fy2025-es.html /
fy2025-pl.html) — there is no .docx source for them, so they can't go
through the same docx -> tagged-PDF export step (docs/generate_report_pdfs.ps1)
that produces the accessible PDF downloads for EN/ZH.

This script closes that gap by reconstructing a properly structured .docx
(headings, table header rows, bold/italic runs, correct document language)
from the existing translated HTML, so all four languages can go through the
same Word "Save As PDF with document structure tags" pipeline.

Usage:
    python docs/generate_report_docx_i18n.py

Re-run after the ES/PL HTML reports change, then re-run
docs/generate_report_pdfs.ps1 to refresh the PDFs.
"""
import os

from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
REPORTS_DIR = os.path.join(HERE, "reports")
FY2025_DIR = os.path.join(ROOT, "FY2025")

LANGS = {
    "es": {
        "html": os.path.join(REPORTS_DIR, "fy2025-es.html"),
        "docx": os.path.join(FY2025_DIR, "FY2025_US_State_HFA_ACFR_Analysis_Report_ES.docx"),
        "lang_tag": "es-ES",
    },
    "pl": {
        "html": os.path.join(REPORTS_DIR, "fy2025-pl.html"),
        "docx": os.path.join(FY2025_DIR, "FY2025_US_State_HFA_ACFR_Analysis_Report_PL.docx"),
        "lang_tag": "pl-PL",
    },
}


def set_default_language(document, lang_tag):
    """Set w:lang on the Normal style's default run properties so every run
    that doesn't override it inherits the document language. This is what
    lets Word's tagged-PDF export stamp the correct /Lang on PDF content
    (required reading for screen readers, and a core PDF/UA checkpoint)."""
    normal = document.styles["Normal"]
    rpr = normal.element.get_or_add_rPr()
    for tag in ("w:lang",):
        existing = rpr.find(qn(tag))
        if existing is not None:
            rpr.remove(existing)
    lang_el = OxmlElement("w:lang")
    lang_el.set(qn("w:val"), lang_tag)
    lang_el.set(qn("w:eastAsia"), lang_tag)
    lang_el.set(qn("w:bidi"), lang_tag)
    rpr.append(lang_el)


def add_inline_runs(paragraph, node, bold=False, italic=False):
    """Recursively walk an HTML node's children, emitting Word runs with
    bold/italic carried down from ancestor <strong>/<em> tags."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                run = paragraph.add_run(text)
                run.bold = bold
                run.italic = italic
        elif child.name == "br":
            paragraph.add_run().add_break()
        elif child.name == "strong":
            add_inline_runs(paragraph, child, bold=True, italic=italic)
        elif child.name == "em":
            add_inline_runs(paragraph, child, bold=bold, italic=True)
        else:
            add_inline_runs(paragraph, child, bold=bold, italic=italic)


def mark_header_row(row):
    """Flag a table row as a repeating header row (w:tblHeader). Word's
    tagged-PDF export turns this into an accessible table header (<TH> /
    Table Header Cell) rather than a plain data cell."""
    trPr = row._tr.get_or_add_trPr()
    header_el = OxmlElement("w:tblHeader")
    trPr.append(header_el)


def build_table(document, table_tag):
    rows = table_tag.find_all("tr", recursive=False)
    if not rows:
        return
    ncols = max(len(r.find_all("td", recursive=False)) for r in rows)
    table = document.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for ridx, tr in enumerate(rows):
        cells = tr.find_all("td", recursive=False)
        for cidx, td in enumerate(cells):
            cell = table.cell(ridx, cidx)
            cell.paragraphs[0].text = ""
            text = td.get_text(strip=True)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            if ridx == 0:
                run.bold = True
        if ridx == 0:
            mark_header_row(table.rows[0])
    # blank paragraph after the table for visual breathing room
    document.add_paragraph()


def convert(lang, cfg):
    with open(cfg["html"], encoding="utf-8") as fh:
        soup = BeautifulSoup(fh.read(), "html.parser")

    document = Document()
    set_default_language(document, cfg["lang_tag"])

    top_level = [el for el in soup.children if getattr(el, "name", None)]

    title_text = None
    for el in top_level:
        if el.name == "h1":
            break
        if el.name == "p" and el.find("strong"):
            title_text = el.get_text(" ", strip=True)
            break

    for el in top_level:
        if el.name in ("h1", "h2", "h3"):
            level = int(el.name[1])
            heading = document.add_heading(level=level)
            add_inline_runs(heading, el)
        elif el.name == "p":
            p = document.add_paragraph()
            add_inline_runs(p, el)
        elif el.name in ("ul", "ol"):
            style = "List Bullet" if el.name == "ul" else "List Number"
            for li in el.find_all("li", recursive=False):
                p = document.add_paragraph(style=style)
                add_inline_runs(p, li)
        elif el.name == "table":
            build_table(document, el)

    cp = document.core_properties
    cp.language = cfg["lang_tag"]
    if title_text:
        cp.title = title_text
    cp.author = "MD-ai-fin"

    os.makedirs(os.path.dirname(cfg["docx"]), exist_ok=True)
    document.save(cfg["docx"])
    print(f"Wrote {cfg['docx']}")


def main():
    for lang, cfg in LANGS.items():
        convert(lang, cfg)


if __name__ == "__main__":
    main()
