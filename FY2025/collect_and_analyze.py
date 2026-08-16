#!/usr/bin/env python3
"""Collect FY2025 state housing authority ACFRs and generate comparative Word summary."""

from __future__ import annotations

import json
import re
import sys
import time
import urllib.parse
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import Optional

import pdfplumber
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

BASE_DIR = Path(__file__).resolve().parent
CATALOG_PATH = BASE_DIR / "state_hfa_catalog.json"
PDF_DIR = BASE_DIR / "pdf"
TXT_DIR = BASE_DIR / "txt"
DATA_PATH = BASE_DIR / "fy2025_metrics.json"
LOG_PATH = BASE_DIR / "download_log.json"
REPORT_PATH = BASE_DIR / "FY2025_US_State_HFA_ACFR_Analysis_Report.docx"
REPORT_PATH_CN = BASE_DIR / "FY2025_美国州级住房金融机构ACFR综合分析报告.docx"
DIRECT_URLS_PATH = BASE_DIR / "direct_urls.json"
DOCS_DIR = BASE_DIR.parent / "docs"

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
    )
}

MANUAL_OVERRIDES: dict[str, dict] = {
    "CA": {
        "net_position_2025": 3_442_918_000,
        "net_position_2024": 3_220_848_000,
        "net_position_change": 222_070_000,
        "net_position_change_pct": 6.9,
        "total_assets_2025": 4_980_031_000,
        "total_assets_2024": 4_558_601_000,
        "total_liabilities_2025": 1_511_418_000,
    },
    "AL": {
        "net_position_2025": 516_018_000,
        "net_position_2024": 484_375_000,
        "total_assets_2025": 1_364_901_000,
        "total_assets_2024": 1_082_234_000,
        "total_liabilities_2025": 846_539_000,
        "total_liabilities_2024": 596_731_000,
    },
    "FL": {
        # ACFR also shows "$ 4,771.1" in millions; extractor wrongly treated that as dollars.
        "net_position_2025": 4_771_085_657,
        "net_position_2024": 4_212_748_551,
        "net_position_change": 558_337_106,
        "net_position_change_pct": 13.25,
        "total_assets_2025": 7_722_418_892,
        "total_liabilities_2025": 2_951_333_235,
    },
    "DE": {"net_position_2025": 732_459_779, "net_position_change_pct": 3.8},
    "ID": {"net_position_2025": 716_007_000, "total_assets_2025": 6_026_206_000, "net_position_change_pct": 8.5},
    "IL": {
        "net_position_2025": 1_533_700_000,
        "net_position_change": 159_900_000,
        "net_position_change_pct": 11.6,
        "total_assets_2025": 7_846_500_000,
    },
    "IN": {
        "net_position_2025": 599_528_335,
        "net_position_2024": 497_849_911,
        "total_assets_2025": 3_247_001_038,
        "total_assets_2024": 2_740_054_627,
        "total_liabilities_2025": 2_646_002_390,
        "total_liabilities_2024": 2_242_119_665,
        "net_position_change_pct": 20.4,
    },
    "LA": {"net_position_2025": 821_730_000, "total_assets_2025": 1_545_038_000, "net_position_change_pct": 31.9},
    "IA": {
        "net_position_2025": 1_664_618_000,
        "net_position_2024": 1_530_373_000,
        "total_assets_2025": 6_173_161_000,
        "total_assets_2024": 5_626_808_000,
        "total_liabilities_2025": 4_487_552_000,
        "total_liabilities_2024": 4_071_413_000,
        "net_position_change": 134_245_000,
        "net_position_change_pct": 8.8,
    },
    "MA": {
        "net_position_2025": 2_083_948_000,
        "net_position_2024": 1_822_847_000,
        "total_assets_2025": 8_079_680_000,
        "total_assets_2024": 7_519_040_000,
        "total_liabilities_2025": 5_998_525_000,
        "total_liabilities_2024": 5_684_807_000,
        "net_position_change_pct": 14.3,
    },
    "MI": {
        "net_position_2025": 1_292_050_000,
        "net_position_2024": 1_160_992_000,
        "total_assets_2025": 8_027_764_000,
        "total_assets_2024": 6_868_705_000,
        "total_liabilities_2025": 6_685_707_000,
        "total_liabilities_2024": 5_662_217_000,
        "net_position_change_pct": 11.3,
    },
    "MN": {"net_position_2025": 1_649_443_000, "total_assets_2025": 8_543_311_000},
    "MO": {
        "net_position_2025": 708_007_000,
        "net_position_2024": 664_589_000,
        "total_assets_2025": 3_192_551_000,
        "total_assets_2024": 2_461_129_000,
        "total_liabilities_2025": 2_491_534_000,
        "total_liabilities_2024": 1_806_271_000,
        "net_position_change_pct": 6.5,
    },
    "CO": {
        "net_position_2025": 971_675_000,
        "net_position_2024": 828_465_000,
        "total_assets_2025": 8_385_309_000,
        "net_position_change_pct": 17.3,
    },
    "NC": {
        "net_position_2025": 979_793_000,
        "net_position_2024": 856_993_000,
        "total_assets_2025": 4_815_950_000,
        "net_position_change_pct": 14.3,
    },
    "ND": {
        "net_position_2025": 280_805_000,
        "net_position_2024": 256_450_000,
        "total_assets_2025": 2_818_952_000,
        "total_assets_2024": 2_255_841_000,
        "total_liabilities_2025": 2_521_600_000,
        "total_liabilities_2024": 1_978_564_000,
        "net_position_change_pct": 9.0,
    },
    "NE": {
        "net_position_2025": 454_474_000,
        "net_position_2024": 432_973_000,
        "total_assets_2025": 2_910_591_000,
        "total_assets_2024": 2_249_850_000,
        "total_liabilities_2025": 2_447_488_000,
        "total_liabilities_2024": 1_805_585_000,
        "net_position_change_pct": 5.0,
    },
    "AR": {
        "net_position_2025": 530_822_000,
        "net_position_2024": 505_820_000,
        "total_assets_2025": 1_142_207_000,
        "total_assets_2024": 1_008_195_000,
        "total_liabilities_2025": 611_156_000,
        "total_liabilities_2024": 502_172_000,
        "net_position_change_pct": 5.1,
    },
    "ME": {
        "net_position_2025": 508_535_000,
        "net_position_2024": 475_829_000,
        "net_position_change_pct": 6.9,
    },
    "UT": {
        "net_position_2025": 610_650_000,
        "net_position_2024": 514_400_000,
        "total_assets_2025": 3_861_330_000,
        "total_assets_2024": 2_923_819_000,
        "net_position_change_pct": 18.7,
    },
    "CT": {
        "net_position_2025": 957_039_000,
        "net_position_2024": 764_506_000,
        "total_assets_2025": 8_122_355_000,
        "total_assets_2024": 7_043_493_000,
        "total_liabilities_2025": 6_986_013_000,
        "total_liabilities_2024": 6_115_720_000,
        "net_position_change_pct": 25.2,
    },
    "KY": {
        "net_position_2025": 519_377_000,
        "net_position_2024": 486_274_000,
        "total_assets_2025": 1_434_509_000,
        "total_assets_2024": 1_180_300_000,
        "total_liabilities_2025": 915_132_000,
        "total_liabilities_2024": 694_000_000,
        "net_position_change_pct": 6.8,
    },
    "SC": {
        "net_position_2025": 612_736_978,
        "net_position_2024": 570_045_294,
        "total_assets_2025": 2_330_498_465,
        "total_assets_2024": 1_847_789_251,
        "total_liabilities_2025": 1_717_292_827,
        "total_liabilities_2024": 1_275_296_701,
        "net_position_change_pct": 7.5,
    },
    "VT": {
        "net_position_2025": 123_644_000,
        "net_position_2024": 112_192_000,
        "total_assets_2025": 679_864_000,
        "total_assets_2024": 565_610_000,
        "total_liabilities_2025": 581_488_000,
        "total_liabilities_2024": 480_222_000,
        "net_position_change_pct": 10.2,
    },
    "OK": {
        "net_position_2025": 417_517_722,
        "net_position_2024": 387_013_564,
        "total_assets_2025": 1_472_100_000,
        "total_assets_2024": 1_091_700_000,
        "total_liabilities_2025": 1_054_900_000,
        "total_liabilities_2024": 707_100_000,
        "net_position_change_pct": 7.9,
    },
    "KS": {
        "net_position_2025": 78_666_000,
        "net_position_2024": 74_066_000,
        "total_assets_2025": 96_859_000,
        "total_liabilities_2025": 20_134_000,
        "net_position_change_pct": 6.2,
    },
    "NM": {
        "net_position_2025": 323_285_000,
        "net_position_2024": 290_095_000,
        "total_assets_2025": 2_716_367_000,
        "total_assets_2024": 2_376_981_000,
        "total_liabilities_2025": 2_393_131_000,
        "total_liabilities_2024": 2_086_795_000,
        "net_position_change_pct": 11.4,
    },
    "OH": {"net_position_2025": 526_058_065, "total_assets_2025": 3_946_798_760, "net_position_change_pct": 19.8},
    "TN": {"net_position_2025": 665_000_000, "net_position_change_pct": 4.74},
    "TX": {
        "net_position_2025": 1_078_974_997,
        "net_position_2024": 970_575_954,
        "total_assets_2025": 4_742_568_191,
        "total_assets_2024": 4_415_168_572,
        "total_liabilities_2025": 3_674_048_947,
        "total_liabilities_2024": 3_443_251_284,
    },
    "VA": {
        "net_position_2025": 3_963_105_718,
        "net_position_2024": 3_859_471_718,
        "net_position_change": 103_634_000,
        "net_position_change_pct": 2.7,
        "total_assets_2025": 12_350_130_739,
        "total_assets_2024": 10_909_544_382,
        "total_liabilities_2025": 8_313_327_538,
        "total_liabilities_2024": 6_984_867_188,
    },
    "WA": {
        "net_position_2025": 1_064_113_600,
        "net_position_2024": 935_580_289,
        "total_assets_2025": 2_525_718_782,
        "total_assets_2024": 2_162_353_345,
        "total_liabilities_2025": 1_456_138_353,
        "total_liabilities_2024": 1_221_592_287,
        "net_position_change_pct": 13.7,
    },
    "WI": {
        "net_position_2025": 1_134_100_000,
        "net_position_2024": 964_200_000,
        "net_position_change_pct": 17.6,
        "total_assets_2025": 4_467_400_000,
        "total_assets_2024": 4_160_100_000,
        "total_liabilities_2025": 3_316_300_000,
        "total_liabilities_2024": 3_174_000_000,
    },
    "WV": {"net_position_2025": 640_049_000, "total_assets_2025": 1_607_042_000, "net_position_change_pct": 4.7},
    "WY": {
        "net_position_2025": 456_281_420,
        "net_position_2024": 441_265_649,
        "total_assets_2025": 1_540_498_559,
        "total_assets_2024": 1_405_305_159,
        "total_liabilities_2025": 1_070_832_241,
        "total_liabilities_2024": 948_673_149,
        "net_position_change_pct": 3.4,
    },
    "HI": {
        "net_position_2025": 2_653_188_757,
        "net_position_2024": 2_330_519_000,
        "total_assets_2025": 2_708_698_081,
        "total_assets_2024": 2_389_133_000,
        "total_liabilities_2025": 52_702_210,
        "net_position_change_pct": 13.8,
    },
    "MT": {
        "net_position_2025": 175_942_443,
        "net_position_2024": 169_439_799,
        "total_assets_2025": 886_009_021,
        "total_assets_2024": 794_690_515,
        "total_liabilities_2025": 710_209_004,
        "total_liabilities_2024": 625_413_802,
        "net_position_change_pct": 3.84,
    },
}

EXCLUDE_FROM_RANKINGS: set[str] = set()  # PA figures verified as PHFA's own entity-level MD&A data, not a state-consolidated total; included in rankings

FY2025_PATTERNS = [
    r"fy\s*2025",
    r"fiscal\s+year\s+2025",
    r"year\s+ended\s+june\s+30,\s*2025",
    r"june\s+30,\s*2025",
    r"6/30/25",
    r"06/30/2025",
    r"2024/2025",
    r"2025\s+acfr",
    r"fy25",
    r"fy-?2025",
]

PDF_HINTS = [
    "acfr",
    "annual comprehensive",
    "financial statement",
    "audited",
    "annual report",
    "annual financial",
]


@dataclass
class AgencyMetrics:
    state: str
    name: str
    abbr: str
    source_file: str = ""
    source_url: str = ""
    status: str = "missing"
    net_position_2025: Optional[float] = None
    net_position_2024: Optional[float] = None
    net_position_change: Optional[float] = None
    net_position_change_pct: Optional[float] = None
    total_assets_2025: Optional[float] = None
    total_assets_2024: Optional[float] = None
    total_liabilities_2025: Optional[float] = None
    total_liabilities_2024: Optional[float] = None
    revenues_2025: Optional[float] = None
    expenses_2025: Optional[float] = None
    net_position_ratio: Optional[float] = None
    notes: list[str] = field(default_factory=list)


def ensure_dirs() -> None:
    PDF_DIR.mkdir(parents=True, exist_ok=True)
    TXT_DIR.mkdir(parents=True, exist_ok=True)


def load_catalog() -> list[dict]:
    with CATALOG_PATH.open(encoding="utf-8") as f:
        catalog = json.load(f)
    if DIRECT_URLS_PATH.exists():
        direct = json.loads(DIRECT_URLS_PATH.read_text(encoding="utf-8"))
        by_state = {item["state"]: item for item in catalog}
        for entry in direct:
            state = entry["state"]
            if state in by_state:
                urls = by_state[state].setdefault("direct_pdfs", [])
                if entry["url"] not in urls:
                    urls.insert(0, entry["url"])
    return catalog


def normalize_url(base: str, href: str) -> str:
    return urllib.parse.urljoin(base, href)


def looks_like_fy2025(text: str) -> bool:
    t = text.lower()
    return any(re.search(p, t) for p in FY2025_PATTERNS)


def looks_like_financial_pdf(text: str) -> bool:
    t = text.lower()
    return any(h in t for h in PDF_HINTS)


def parse_money(raw: str) -> Optional[float]:
    if not raw:
        return None
    s = raw.strip().replace("$", "").replace(",", "").replace(" ", "")
    if s in {"", "-", "—", "N/A", "n/a"}:
        return None
    if s.startswith("(") and s.endswith(")"):
        s = "-" + s[1:-1]
    try:
        return float(s)
    except ValueError:
        return None


def extract_text_from_pdf(pdf_path: Path) -> str:
    chunks: list[str] = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            chunks.append(text)
    return "\n".join(chunks)


def first_match(patterns: list[str], text: str, flags=re.I | re.S) -> Optional[re.Match]:
    for pat in patterns:
        m = re.search(pat, text, flags)
        if m:
            return m
    return None


def normalize_financial_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    text = re.sub(r"(\d)\s+(?=\d)", r"\1", text)
    text = re.sub(r"\s+", " ", text)
    return text


def to_dollars(value: Optional[float], unit: str) -> Optional[float]:
    if value is None:
        return None
    unit = unit.lower()
    if unit == "thousand":
        return value * 1_000
    if unit == "million":
        return value * 1_000_000
    if unit == "billion":
        return value * 1_000_000_000
    return value


def plausible_np(value: Optional[float]) -> bool:
    return value is not None and 1_000_000 <= value <= 25_000_000_000


def plausible_assets(value: Optional[float]) -> bool:
    return value is not None and 10_000_000 <= value <= 100_000_000_000


def detect_unit_scale(text: str) -> str:
    head = text[:12000].lower()
    if "in thousands of dollars" in head or "(in thousands)" in head:
        return "thousand"
    if "millions of dollars" in head or "in millions" in head:
        return "million"
    return "raw"


def extract_metrics(text: str, agency: dict) -> AgencyMetrics:
    m = AgencyMetrics(state=agency["state"], name=agency["name"], abbr=agency["abbr"])
    t = normalize_financial_text(text)
    scale = detect_unit_scale(text)

    hl_patterns = [
        (
            r"overall net position increased by \$(?P<chg>[\d,\.]+) million, to \$(?P<np>[\d,\.]+) million",
            "million",
            "million",
        ),
        (
            r"Total net position as of June 30, 2025, was \$(?P<np>[\d,\.]+) million, an increase of \$(?P<chg>[\d,\.]+) million or (?P<pct>[\d,\.]+)%",
            "million",
            "million",
        ),
        (
            r"net position increased \$(?P<chg>[\d,\.]+) million to \$(?P<np>[\d,\.]+) million",
            "million",
            "million",
        ),
        (
            r"Net position increased \$(?P<chg>[\d,\.]+) million from .*? to \$(?P<np>[\d,\.]+) million",
            "million",
            "million",
        ),
        (
            r"net position totaled \$(?P<np>[\d,\.]+) billion at June 30, 2025",
            "billion",
            None,
        ),
        (
            r"Total Net Position improved by \$(?P<chg>[\d,\.]+).*?to \$(?P<np>[\d,\.]+).*?June 30,\s*2025",
            "raw",
            "raw",
        ),
        (
            r"change in net position of \$(?P<chg>[\d,\.]+) million for the fiscal year ended June 30, 2025",
            None,
            "million",
        ),
        (
            r"Net Position \$ (?P<np>[\d,\.]+)",
            "thousand",
            None,
        ),
    ]
    for pattern, np_unit, chg_unit in hl_patterns:
        hit = re.search(pattern, t, re.I)
        if not hit:
            continue
        if hit.groupdict().get("np") and not m.net_position_2025:
            np_val = to_dollars(parse_money(hit.group("np")), np_unit or scale)
            if plausible_np(np_val):
                m.net_position_2025 = np_val
        if hit.groupdict().get("chg") and not m.net_position_change:
            chg_val = to_dollars(parse_money(hit.group("chg")), chg_unit or scale)
            if chg_val is not None:
                m.net_position_change = chg_val
        if hit.groupdict().get("pct") and not m.net_position_change_pct:
            m.net_position_change_pct = parse_money(hit.group("pct"))

    table_patterns = [
        r"Total assets (?P<a25>[\d,\.]+) (?P<a24>[\d,\.]+).*?(?P<pct1>[\d,\.]+)%",
        r"Total net position (?P<np25>[\d,\.]+) (?P<np24>[\d,\.]+).*?(?P<pct2>[\d,\.]+)%",
    ]
    if scale == "thousand":
        for pat in table_patterns:
            hit = re.search(pat, t, re.I)
            if not hit:
                continue
            gd = hit.groupdict()
            if gd.get("np25") and not m.net_position_2025:
                np25 = to_dollars(parse_money(gd["np25"]), "thousand")
                if plausible_np(np25):
                    m.net_position_2025 = np25
            if gd.get("np24") and not m.net_position_2024:
                np24 = to_dollars(parse_money(gd["np24"]), "thousand")
                if plausible_np(np24):
                    m.net_position_2024 = np24
            if gd.get("a25") and not m.total_assets_2025:
                a25 = to_dollars(parse_money(gd["a25"]), "thousand")
                if plausible_assets(a25):
                    m.total_assets_2025 = a25
            if gd.get("a24") and not m.total_assets_2024:
                a24 = to_dollars(parse_money(gd["a24"]), "thousand")
                if plausible_assets(a24):
                    m.total_assets_2024 = a24
            if gd.get("pct2") and not m.net_position_change_pct:
                m.net_position_change_pct = parse_money(gd["pct2"])

    if m.net_position_2025 and m.net_position_2024 and not m.net_position_change:
        m.net_position_change = m.net_position_2025 - m.net_position_2024
    if m.net_position_change and m.net_position_2024 and m.net_position_2024 != 0 and not m.net_position_change_pct:
        m.net_position_change_pct = (m.net_position_change / m.net_position_2024) * 100
    if m.net_position_2025 and m.total_assets_2025 and m.total_assets_2025 > 0:
        m.net_position_ratio = (m.net_position_2025 / m.total_assets_2025) * 100

    if not any([m.net_position_2025, m.total_assets_2025]):
        m.notes.append("未能自动提取核心指标，需人工复核PDF")

    overrides = MANUAL_OVERRIDES.get(agency["state"], {})
    for key, val in overrides.items():
        setattr(m, key, val)
    if agency["state"] in EXCLUDE_FROM_RANKINGS:
        m.notes.append("文件为州级ACFR/合并报表，未纳入HFA专项排名")
        m.net_position_2025 = overrides.get("net_position_2025")
        m.total_assets_2025 = overrides.get("total_assets_2025")
        m.net_position_change_pct = overrides.get("net_position_change_pct")
    if m.net_position_2025 and m.total_assets_2025 and m.total_assets_2025 > 0:
        m.net_position_ratio = (m.net_position_2025 / m.total_assets_2025) * 100
    return m


def discover_pdf_links(page_url: str) -> list[str]:
    try:
        resp = requests.get(page_url, headers=HEADERS, timeout=30)
        resp.raise_for_status()
    except Exception:
        return []

    soup = BeautifulSoup(resp.text, "html.parser")
    links: list[str] = []
    for a in soup.find_all("a", href=True):
        href = a["href"]
        if not href.lower().endswith(".pdf"):
            continue
        label = (a.get_text(" ", strip=True) + " " + href).strip()
        if looks_like_financial_pdf(label) and looks_like_fy2025(label):
            links.append(normalize_url(page_url, href))
    return links


def download_pdf(url: str, dest: Path, retries: int = 3) -> bool:
    last_error = None
    for attempt in range(retries):
        for verify in (True, False):
            try:
                resp = requests.get(
                    url,
                    headers=HEADERS,
                    timeout=90,
                    allow_redirects=True,
                    verify=verify,
                )
                resp.raise_for_status()
                content = resp.content
                if content.startswith(b"%PDF") and len(content) >= 5000:
                    dest.write_bytes(content)
                    return True
            except Exception as exc:
                last_error = exc
        time.sleep(0.8)
    return False


def collect_agency(agency: dict, log: dict) -> Optional[Path]:
    state = agency["state"]
    abbr = agency["abbr"]
    candidates: list[str] = list(agency.get("direct_pdfs", []))

    for page in agency.get("urls", []):
        candidates.extend(discover_pdf_links(page))
        time.sleep(0.4)

    # dedupe preserve order
    seen = set()
    unique_candidates = []
    for u in candidates:
        if u not in seen:
            seen.add(u)
            unique_candidates.append(u)

    log[state] = {"candidates": unique_candidates, "downloaded": None, "errors": []}

    for url in unique_candidates:
        filename = f"{state}_{abbr}_FY2025_ACFR.pdf"
        dest = PDF_DIR / filename
        if dest.exists() and dest.stat().st_size > 5000:
            log[state]["downloaded"] = str(dest)
            return dest
        ok = download_pdf(url, dest)
        if ok and dest.stat().st_size > 5000:
            log[state]["downloaded"] = str(dest)
            log[state]["source_url"] = url
            return dest
        if dest.exists():
            dest.unlink(missing_ok=True)
        log[state]["errors"].append(f"failed: {url}")
        time.sleep(0.3)

    return None


def fmt_money(val: Optional[float], unit: str = "auto") -> str:
    if val is None:
        return "N/A"
    if unit == "auto":
        if abs(val) >= 1_000_000:
            return f"${val/1_000_000:,.1f}M"
        if abs(val) >= 1_000:
            return f"${val/1_000:,.1f}K"
        return f"${val:,.0f}"
    if unit == "M":
        return f"${val/1_000_000:,.1f}M"
    if unit == "B":
        return f"${val/1_000_000_000:,.2f}B"
    return f"${val:,.0f}"


def fmt_pct(val: Optional[float]) -> str:
    if val is None:
        return "N/A"
    return f"{val:+.1f}%"


# ---- Enrichment data (population, program activity, data quality) ----

# Mirrors PROGRAM_ACTIVITY_METRIC_BUCKETS in docs/program_activity.js so the
# report aggregates match the web map's right-hand panel exactly. Order matters:
# rental_assistance is matched before the broad multifamily/rental bucket.
PROGRAM_ACTIVITY_BUCKETS: list[tuple] = [
    ("pa_homeownership", "pa_homeownership_households", "pa_homeownership_amount",
     re.compile(r"homeown|homebuyer|single_family|downpayment|covenant|sonyma|mortgage_operations")),
    ("pa_repair", "pa_repair_households", "pa_repair_amount",
     re.compile(r"repair|renew|rehab|counsel")),
    ("pa_covid", "pa_covid_households", "pa_covid_amount",
     re.compile(r"covid|emergency|foreclosure|homeless")),
    ("pa_rental", "pa_rental_units", "pa_rental_amount",
     re.compile(r"rental_assistance|community_affairs|voucher|public_housing|section_?8|tbra|housing_stability")),
    ("pa_multifamily", "pa_multifamily_units", "pa_multifamily_amount",
     re.compile(r"multifamily|rental|units_completed|placed_in_service|build_for|neighborhood|housing_solutions|lihtc|tax_credit|credit_award|hmmf|lmir|legislative_loan")),
    ("pa_energy", "pa_energy_households", "pa_energy_amount",
     re.compile(r"weatheriz|energy|liheap")),
    ("pa_grants", "pa_grants_units", "pa_grants_amount",
     re.compile(r"trust_fund|phare|grant|reach_virginia|federal_assistance|home_program")),
]


def _num(v):
    if v is None or v == "":
        return None
    try:
        return float(v)
    except (TypeError, ValueError):
        return None


def _bucket_metric_for_category(key: str) -> Optional[str]:
    """First-match-wins bucket lookup, mirroring program_activity.js's
    bucketMetricForCategory so the report aggregates match the web map exactly."""
    for metric, _units_key, _amount_key, regex in PROGRAM_ACTIVITY_BUCKETS:
        if regex.search(key):
            return metric
    return None


def load_population_data() -> dict:
    """Return {state: state_record} from docs/state_data.json (2025 metrics)."""
    path = DOCS_DIR / "state_data.json"
    if not path.exists():
        return {}
    data = json.loads(path.read_text(encoding="utf-8"))
    return {s.get("state"): s for s in data.get("states", []) if s.get("state")}


def load_program_activity_aggregates() -> dict:
    """Aggregate the 7 housing-activity buckets across docs/*_program_activity.json."""
    result: dict = {}
    if not DOCS_DIR.exists():
        return result
    for f in sorted(DOCS_DIR.glob("*_program_activity.json")):
        try:
            raw = json.loads(f.read_text(encoding="utf-8"))
        except Exception:
            continue
        if not isinstance(raw, dict):
            continue
        for abbr, body in raw.items():
            if not isinstance(body, dict):
                continue
            cats = body.get("categories") or []
            out: dict = {}
            for metric, units_key, amount_key, regex in PROGRAM_ACTIVITY_BUCKETS:
                matched = [
                    c for c in cats
                    if isinstance(c, dict) and _bucket_metric_for_category(c.get("key") or "") == metric
                ]
                units = None
                amount = None
                for c in matched:
                    q = _num(c.get("fy2025_actual"))
                    if q is not None:
                        units = (units or 0.0) + q
                    cat_amt = 0.0
                    has_amt = False
                    for a in c.get("fy2025_amounts") or []:
                        av = _num(a.get("amount")) if isinstance(a, dict) else None
                        if av is not None:
                            cat_amt += av
                            has_amt = True
                    if has_amt:
                        amount = (amount or 0.0) + cat_amt
                out[units_key] = units
                out[amount_key] = amount
            result[abbr] = out
    return result


def _fmt_int(v: Optional[float]) -> str:
    return f"{int(v):,}" if v is not None else "—"


def _fmt_usd(v: Optional[float]) -> str:
    if v is None:
        return "—"
    if abs(v) >= 1e9:
        return f"${v/1e9:.2f}B"
    if abs(v) >= 1e6:
        return f"${v/1e6:.1f}M"
    return f"${v:,.0f}"


def _enrichment_numbers(metrics: list, language: str) -> dict:
    """Compute formatted figures for the population / housing-activity / data-quality
    sections. Returns a dict of pre-formatted strings plus a `housing_table` list."""
    sep = ", " if language == "en" else "、"
    op, cp = ("(", ")") if language == "en" else ("（", "）")

    def sv(state: str, val: str) -> str:
        return f"{state}{op}{val}{cp}"

    pop_data = load_population_data()
    pa = load_program_activity_aggregates()

    def y2025(rec, key):
        y = rec.get("2025") or {}
        return y.get(key)

    states = list(pop_data.items())

    def ranked(key):
        pairs = [(st, y2025(rec, key)) for st, rec in states if y2025(rec, key) is not None]
        return sorted(pairs, key=lambda x: x[1] or 0, reverse=True)

    # --- population ---
    pop_sorted = ranked("population")
    pop_total = f"{sum(v for _, v in pop_sorted if v)/1e6:.1f}M"
    pop_top = sep.join(sv(st, f"{v/1e6:.1f}M") for st, v in pop_sorted[:5])
    pop_bottom = sep.join(sv(st, f"{v/1e6:.1f}M") for st, v in pop_sorted[-5:])

    inc_sorted = ranked("personal_income_per_capita")
    income_top = sep.join(sv(st, f"${int(v):,}") for st, v in inc_sorted[:3])
    income_bottom = sep.join(sv(st, f"${int(v):,}") for st, v in inc_sorted[-3:])
    income_ratio = f"{inc_sorted[0][1]/inc_sorted[-1][1]:.1f}"

    unemp = [v for st, rec in states if (v := y2025(rec, "unemployment_rate")) is not None]
    unemp_sorted = ranked("unemployment_rate")
    unemp_avg = f"{sum(unemp)/len(unemp):.1f}%"
    unemp_high = sv(unemp_sorted[0][0], f"{unemp_sorted[0][1]:.1f}%")
    unemp_low = sv(unemp_sorted[-1][0], f"{unemp_sorted[-1][1]:.1f}%")

    pov = [v for st, rec in states if (v := y2025(rec, "poverty_rate")) is not None]
    pov_sorted = ranked("poverty_rate")
    pov_avg = f"{sum(pov)/len(pov):.1f}%"
    pov_high = sep.join(sv(st, f"{v:.1f}%") for st, v in pov_sorted[:3])
    pov_low = sv(pov_sorted[-1][0], f"{pov_sorted[-1][1]:.1f}%")

    price_sorted = ranked("median_home_price")
    price_high = sv(price_sorted[0][0], f"${int(price_sorted[0][1]):,}")
    price_low = sv(price_sorted[-1][0], f"${int(price_sorted[-1][1]):,}")
    price_ratio = f"{price_sorted[0][1]/price_sorted[-1][1]:.1f}"

    homeless_sorted = ranked("homeless_count")
    homeless_total = _fmt_int(sum(v for _, v in homeless_sorted if v))
    homeless_top = sep.join(sv(st, f"{int(v):,}") for st, v in homeless_sorted[:2])
    homeless_low = sv(homeless_sorted[-1][0], f"{int(homeless_sorted[-1][1]):,}")

    # --- program activity ---
    pa_states = str(len(pa))

    def pa_count(key):
        return sum(1 for v in pa.values() if v.get(key) is not None)

    def pa_sum(key):
        return sum(v.get(key) or 0 for v in pa.values())

    def pa_top(key, n=3):
        ranked_pairs = sorted(
            [(st, v[key]) for st, v in pa.items() if v.get(key) is not None],
            key=lambda x: x[1], reverse=True,
        )
        return sep.join(sv(st, f"{int(val):,}") for st, val in ranked_pairs[:n])

    mf_amt = _fmt_usd(pa_sum("pa_multifamily_amount"))
    ho_amt = _fmt_usd(pa_sum("pa_homeownership_amount"))
    mf_top = pa_top("pa_multifamily_units", 3)
    ho_top = pa_top("pa_homeownership_households", 3)
    ra_units = _fmt_int(pa_sum("pa_rental_units"))
    ra_states = str(pa_count("pa_rental_units"))
    ra_amt = _fmt_usd(pa_sum("pa_rental_amount"))
    ra_top = pa_top("pa_rental_units", 2)
    en_units = _fmt_int(pa_sum("pa_energy_households"))
    en_top = pa_top("pa_energy_households", 2)
    en_states = str(pa_count("pa_energy_households"))
    repair_states = str(pa_count("pa_repair_households"))
    covid_states = str(pa_count("pa_covid_households"))
    pa_hi = str(pa_count("pa_multifamily_units"))
    pa_lo = str(pa_count("pa_repair_households"))

    housing_table = []
    for metric, units_key, amount_key, _ in PROGRAM_ACTIVITY_BUCKETS:
        housing_table.append({
            "n_states": pa_count(units_key),
            "units": _fmt_int(pa_sum(units_key)),
            "amount": _fmt_usd(pa_sum(amount_key)),
        })

    # --- data quality ---
    downloaded = [m for m in metrics if m.status == "downloaded"]
    fin_states = str(len(downloaded))
    fin_pct = f"{len(downloaded)/len(metrics)*100:.0f}%"
    fin_missing = sep.join(m.state for m in metrics if m.status != "downloaded")

    return {
        "pop_total": pop_total, "pop_top": pop_top, "pop_bottom": pop_bottom,
        "income_top": income_top, "income_bottom": income_bottom, "income_ratio": income_ratio,
        "unemp_avg": unemp_avg, "unemp_high": unemp_high, "unemp_low": unemp_low,
        "pov_avg": pov_avg, "pov_high": pov_high, "pov_low": pov_low,
        "price_high": price_high, "price_low": price_low, "price_ratio": price_ratio,
        "homeless_total": homeless_total, "homeless_top": homeless_top, "homeless_low": homeless_low,
        "pa_states": pa_states, "mf_amt": mf_amt, "ho_amt": ho_amt,
        "mf_top": mf_top, "ho_top": ho_top, "ra_units": ra_units, "ra_states": ra_states,
        "ra_amt": ra_amt, "ra_top": ra_top, "en_units": en_units, "en_top": en_top,
        "en_states": en_states, "repair_states": repair_states, "covid_states": covid_states,
        "pa_hi": pa_hi, "pa_lo": pa_lo,
        "fin_states": fin_states, "fin_pct": fin_pct, "fin_missing": fin_missing,
        "housing_table": housing_table,
    }


def _set_doc_fonts(doc: Document, language: str) -> None:
    """Pin a single East Asian font on the base styles so CJK text renders with
    one consistent font — and one that has a real bold face (Microsoft YaHei).

    Without an explicit ``w:eastAsia``, Word resolves each run's CJK glyphs to
    whatever default East Asian font the machine is set to (on this box that is
    MS Mincho / MS-Gothic), which has no matching bold face, so Word applies an
    uneven synthetic bold. The result is the messy, inconsistent bold the user
    saw in the exported Chinese PDF."""
    doc._hfa_lang = language  # read by add_table to set per-cell fonts
    if language != "zh":
        return
    east = "微软雅黑"  # Microsoft YaHei — ships with a true Bold face
    latin = "Calibri"
    for style_name in ("Normal", "Heading 1", "Heading 2", "Title",
                       "List Bullet", "List Number", "Normal Table"):
        try:
            st = doc.styles[style_name]
        except KeyError:
            continue
        rpr = st.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
        rfonts.set(qn("w:eastAsia"), east)
        # Drop the theme-based font references (e.g. w:eastAsiaTheme), which Word
        # otherwise resolves in favour of the machine's theme majorEastAsia font
        # (here: MS-Gothic) instead of our explicit w:eastAsia above.
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            full = qn("w:" + attr)
            if full in rfonts.attrib:
                del rfonts.attrib[full]


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _apply_run_font(run, language: str) -> None:
    """Explicitly set the Latin + East Asian fonts on a single run.

    Table cells inherit their font from the table style ("Light Grid Accent 1")
    rather than from Normal, and that table style has no East Asian font
    defined, so CJK cell text fell back to MS-Gothic. Setting the font on each
    run directly sidesteps that fallback."""
    if language != "zh":
        return
    run.font.name = "Calibri"  # sets w:ascii + w:hAnsi
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "微软雅黑")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    language = getattr(doc, "_hfa_lang", "en")
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        _apply_run_font(run, language)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            _apply_run_font(run, language)


def _report_copy(language: str) -> dict:
    downloaded_label = "Downloaded" if language == "en" else "已获取"
    missing_label = "Not available" if language == "en" else "未获取"
    date_fmt = "%B %d, %Y" if language == "en" else "%Y年%m月%d日"

    if language == "en":
        return {
            "date_fmt": date_fmt,
            "downloaded_label": downloaded_label,
            "missing_label": missing_label,
            "title": "FY2025 U.S. State Housing Finance Agencies\nACFR and Financial Statement Analysis",
            "subtitle": (
                "Prepared: {date}  |  Coverage: {total} states/territories  |  "
                "Reports obtained: {downloaded} FY2025 filings"
            ),
            "s1": "I. Executive Summary",
            "s1_body": (
                "Fiscal year 2025 (year ended June 30, 2025 for most agencies; a few report on a calendar year) "
                "was a consequential period for State Housing Finance Agencies (HFAs). Interest rates eased, "
                "competition for Low-Income Housing Tax Credits (LIHTC) intensified, and first-time homebuyer "
                "programs continued to expand. This report compiles and compares publicly available ACFRs and "
                "audited financial statements for HFAs across all 51 states and the District of Columbia. "
                "As of the preparation date, {downloaded} FY2025 reports were obtained "
                "({coverage:.0f}% coverage)."
            ),
            "key_findings": (
                "Among {count} agencies with extracted metrics, aggregate net position is approximately {np} "
                "and aggregate total assets approximately {assets}. These aggregates are indicative of sector "
                "scale only and are not comparable to consolidated financial statements given differing "
                "reporting bases. State HFAs function largely as bond-financed intermediaries: mortgage and "
                "development loans on the asset side, and housing bonds on the liability side. The net "
                "position ratio (net position divided by total assets) is a useful indicator of financial "
                "cushion and resilience."
            ),
            "growth_champion": (
                "The strongest year-over-year increase in net position was recorded by {state} ({abbr}), "
                "at {pct}, bringing net position to {np}. Such gains often coincide with higher bond issuance, "
                "portfolio growth, or fair-value recovery, and should be interpreted alongside housing production "
                "measures (for example, units financed)."
            ),
            "s2": "II. Rankings",
            "s2_1": "2.1 Top 10 by Net Position",
            "s2_1_insight": (
                "Florida, Virginia, and California lead by net position, followed by Hawaii and Massachusetts. "
                "Larger agencies typically operate both multifamily housing revenue bond and single-family "
                "mortgage programs. Size is not a performance scorecard, but it often correlates with stronger "
                "market access, lower marginal funding costs, and greater capacity to finance large affordable "
                "housing transactions when federal policy shifts."
            ),
            "s2_2": "2.2 Top 10 by Net Position Growth",
            "s2_2_insight": (
                "Louisiana recorded the highest year-over-year net position growth, with Connecticut, Indiana, "
                "Ohio, Minnesota, and Utah also posting gains in the mid-teens to mid-20s. Rapid growth commonly "
                "reflects a rebound in bond issuance, releases of loan-loss reserves or fair-value gains, or "
                "stronger fee and spread income from state homebuyer programs. Gains driven mainly by fair-value "
                "changes rather than operating results may prove less durable."
            ),
            "s2_2_empty": (
                "Some reports do not disclose comparable two-year net position growth; this section will "
                "expand as coverage improves."
            ),
            "s2_3": "2.3 Top 10 by Total Assets",
            "s2_4": "2.4 Top 10 by Net Position Ratio",
            "s2_4_insight": (
                "Among larger agencies, Hawaii and California show the highest net position ratios; Florida "
                "and Louisiana also sit well above typical peer levels. Smaller HFAs can post elevated ratios "
                "because of simpler balance sheets, so size and ratio should be read together. A thicker "
                "cushion generally provides more flexibility under asset impairments, interest-rate stress, "
                "or narrower refinancing windows."
            ),
            "s3": "III. Operating Models: Five Archetypes",
            "s3_intro": "FY2025 disclosures suggest five broad operating models useful for cross-state comparison:",
            "prototypes": [
                (
                    "Bond-financed large agencies",
                    "Examples include New York, Virginia, Minnesota, Colorado, Connecticut, Michigan, and "
                    "Illinois: large asset bases with liabilities expanding in parallel, and heavy reliance "
                    "on multifamily and single-family bond programs and GSE channels. New York State HFA is "
                    "the largest by a wide margin, reporting FY2025 total assets of about $21.31 billion "
                    "and total liabilities of about $19.57 billion -- roughly 1.7 times the total assets of "
                    "Virginia, the next-largest agency in this group.",
                ),
                (
                    "High-growth mid-to-large agencies",
                    "Examples include Louisiana, Connecticut, Indiana, Ohio, and Utah: double-digit net "
                    "position growth on mid-to-large balance sheets, illustrating how issuance cycles and "
                    "program mix can strengthen equity.",
                ),
                (
                    "Integrated state-department models",
                    "Texas TDHCA is representative: governmental and business-type funds operate in parallel, "
                    "so the disclosure basis and consolidation scope require careful reading.",
                ),
                (
                    "High-cost coastal markets",
                    "California and Washington are illustrative. Results track home prices, construction "
                    "costs, and project pipelines closely. In FY2025, CalHFA reported net position of about "
                    "$3.44 billion and total assets near $5.0 billion, remaining among the largest state "
                    "housing finance agencies.",
                ),
                (
                    "Resource and low-density states",
                    "Wyoming, North Dakota, South Dakota, and New Mexico often show more volatile issuance "
                    "cycles; individual projects can move the financial statements materially.",
                ),
            ],
            "s4": "IV. Interpreting the Numbers",
            "insights": [
                (
                    "Scale does not guarantee safety",
                    "The largest asset bases often carry the highest liabilities and derivative exposure. "
                    "While interest-rate paths remain uncertain, net position ratio and cash-flow capacity "
                    "deserve as much attention as asset rankings. New York illustrates the point: it is now "
                    "the largest agency by total assets, yet its net position ratio is only about 8%, "
                    "thinner than most of its top-10-by-assets peers, which otherwise range roughly from "
                    "16% to 98%.",
                ),
                (
                    "Net position growth is not housing production",
                    "Under GAAP, net position can change with investment fair value, loss provisions, or bond "
                    "premium amortization. Housing output should be validated against program reports "
                    "(units financed, bonds issued).",
                ),
                (
                    "Disclosure quality has market value",
                    "Agencies that earn the GFOA Certificate of Achievement (for example, Illinois IHDA) "
                    "tend to publish clearer MD&A, which can reduce investor due-diligence costs and, "
                    "indirectly, support tighter bond spreads.",
                ),
            ],
            "s_pop": "V. Population and Socio-Economic Context",
            "s_pop_intro": (
                "Total population across the 51 states and the District of Columbia is approximately {pop_total}. "
                "The five most populous are {pop_top}; the five least populous are {pop_bottom}. Population "
                "figures come from the Census Bureau's population estimates program and cover all 51 "
                "states/territories (100% coverage), making this the most consistent of the three data layers."
            ),
            "s_pop_income": (
                "Per-capita personal income (BEA SAINC4) is highest in the District of Columbia (about "
                "{income_top}) and lowest in Mississippi (about {income_bottom}), roughly {income_ratio} times "
                "lower. High incomes concentrate in the District of Columbia, New England, the Mid-Atlantic, "
                "and the West Coast."
            ),
            "s_pop_labor": (
                "The national average unemployment rate is about {unemp_avg}: Nevada is highest ({unemp_high}) "
                "and South Dakota lowest ({unemp_low}). The average poverty rate is about {pov_avg}: Louisiana, "
                "Mississippi, and New Mexico are highest ({pov_high}), while New Hampshire is lowest "
                "({pov_low}). Mississippi, West Virginia, Alabama, Kentucky, and New Mexico generally show both "
                "lower per-capita income and higher poverty."
            ),
            "s_pop_housing": (
                "Median home prices range from about {price_low} in West Virginia to about {price_high} in "
                "Hawaii (roughly {price_ratio} times). The HUD January 2025 point-in-time count totals about "
                "{homeless_total} people: California is highest ({homeless_top}) and Wyoming lowest "
                "({homeless_low}). High-cost coastal markets (California, Hawaii, the District, Massachusetts, "
                "Washington) combine the highest prices with the heaviest homelessness burden."
            ),
            "s_pop_insight": (
                "Population size correlates broadly with HFA balance-sheet size, but the distribution of "
                "per-capita income, home prices, and homelessness reveals two very different demand structures: "
                "high-cost coastal states face an undersupply-plus-high-price problem, while lower-income "
                "Southern states face an affordability-plus-concentrated-poverty problem. These two pressure "
                "profiles call for different HFA product mixes."
            ),
            "s_housing": "VI. Housing Program Activity Analysis",
            "s_housing_intro": (
                "{pa_states} states and territories have FY2025 housing program-activity modules, compiled "
                "state-by-state from each HFA's annual report program disclosures. The table below aggregates "
                "seven business categories, reporting for each the number of disclosing states, volume "
                "(units/households), and dollar amount to support cross-state comparison."
            ),
            "s_housing_table": ["Business Category", "Disclosing States", "Volume (units/households)", "Amount"],
            "pa_labels": [
                "Homeownership loans", "Repair/counseling", "COVID emergency", "Rental assistance",
                "Multifamily new/preserved", "Energy assistance", "Grants/trust fund",
            ],
            "s_housing_insights": [
                (
                    "Multifamily and homeownership lending dominate capital deployment",
                    "Multifamily new/preserved and homeownership lending together account for about {mf_amt} and "
                    "{ho_amt}. Multifamily volume is led by Texas ({mf_top}); homeownership households by "
                    "Washington ({ho_top}). These two categories form the core of most HFAs' mortgage and "
                    "development-loan assets."
                ),
                (
                    "Rental assistance reaches the most households at the lowest unit cost",
                    "Rental assistance totals about {ra_units} households across {ra_states} states, but only "
                    "about {ra_amt} in amount, reflecting its nature as a recurring small subsidy rather than "
                    "one-time capital investment; Texas ({ra_top}) leads in scale."
                ),
                (
                    "Several categories have low disclosure coverage, so cross-state comparison needs care",
                    "Energy assistance ({en_units} households, concentrated in {en_top} among {en_states} states) "
                    "and grants/trust-fund programs carry meaningful amounts from few states; repair/counseling "
                    "({repair_states} states) and COVID emergency ({covid_states} states) have the lowest "
                    "coverage. Because disclosure conventions differ widely, cross-state comparisons should be "
                    "read as directional rather than precise."
                ),
            ],
            "s_quality": "VII. Data Source and Quality Assessment",
            "s_quality_intro": (
                "This report rests on three data layers with different sources and quality gradients, which "
                "directly affect the robustness of its conclusions."
            ),
            "s_quality_insights": [
                (
                    "Financial data (high confidence, ~90% coverage)",
                    "Net position, total assets, and liabilities come from state ACFRs or audited financial "
                    "statements; {fin_states} states ({fin_pct}) were obtained. {fin_missing} are not included "
                    "because audits are incomplete or only FY2024 filings were posted. Certain core figures were "
                    "manually verified and are noted state-by-state in the coverage inventory."
                ),
                (
                    "Population and socio-economic data (highest confidence, 100% coverage)",
                    "Population, income, unemployment, poverty, home prices, and homelessness all come from "
                    "federal statistics (Census/BEA/BLS/ACS/HUD), covering all 51 states/territories with "
                    "uniform definitions and direct cross-state comparability — the highest-quality layer of "
                    "the three."
                ),
                (
                    "Housing program-activity data (medium confidence, ~86% coverage)",
                    "Program activity comes from each HFA's annual-report program disclosures; {pa_states} states "
                    "have modules, but coverage across the seven categories ranges from {pa_hi} states "
                    "(multifamily) to {pa_lo} states (repair). Unit definitions, amount conventions, and whether "
                    "tax credits are included vary widely, making this the layer that most requires careful "
                    "interpretation."
                ),
                (
                    "Conclusion",
                    "The quality gradient (population/economy ≈ financial > housing activity) determines how "
                    "robust each conclusion is: population and financial insights are relatively reliable, while "
                    "housing-activity insights should be treated as directional rather than precise comparisons."
                ),
            ],
            "s5": "VIII. Recommendations for FY2026 and Beyond",
            "recommendations": [
                (
                    "Adopt a common HFA financial health framework",
                    "Standardize disclosure of net position ratio, debt service coverage, delinquency rates, "
                    "LIHTC pipeline, and the single-family versus multifamily mix so that fragmented narratives "
                    "can be replaced with comparable metrics.",
                ),
                (
                    "Connect affordability outcomes to the financial statements",
                    "Add MD&A unit-cost metrics—affordable units per $1 million of bond proceeds, and "
                    "first-time buyers supported per dollar of net position—to link reported finances to "
                    "housing results.",
                ),
                (
                    "Expand regional bond pooling and risk sharing",
                    "Smaller HFAs may lower issuance costs through regional bond pools. Larger HFAs may "
                    "benefit from more standardized hedging for disaster and interest-rate risk.",
                ),
                (
                    "Publish machine-readable disclosures",
                    "XBRL or structured CSV supplements would help researchers, rating agencies, and federal "
                    "partners monitor HFAs on a more timely basis.",
                ),
                (
                    "Track the missing-middle segment (80%–120% AMI)",
                    "Treat middle-income projects as a distinct planning category so resources are not "
                    "concentrated only at the lowest AMI tiers and the market-rate end of the spectrum, "
                    "which can leave teachers, nurses, and similar workers underserved.",
                ),
            ],
            "s6": "IX. Data Coverage and File Inventory",
            "s7": "X. States Without FY2025 Reports",
            "s7_body": (
                "As of this report, the following {count} states or territories had no publicly locatable "
                "FY2025 ACFR or audited financial statement. Common reasons include incomplete audits, "
                "FY2024-only postings, or filings hosted on nonstandard sites."
            ),
            "appendix": "Appendix: Methodology",
            "appendix_body": (
                "Sources are FY2025 ACFRs or audited financial statement PDFs published by state HFAs or "
                "state auditors. Extraction prioritizes MD&A financial highlights and Statement of Net "
                "Position lines for total assets, total liabilities, and total net position (often reported "
                "in thousands of dollars). Disclosure bases differ across governmental activities, "
                "business-type activities, and blended presentations; cross-state comparisons should be read "
                "with the accompanying notes."
            ),
            "table_rank": ["Rank", "State", "Agency", "FY2025 Net Position", "FY2025 Total Assets", "Net Position Ratio"],
            "table_growth": ["Rank", "State", "Agency", "Net Position Change", "Growth Rate", "FY2025 Net Position"],
            "table_assets": ["Rank", "State", "Agency", "FY2025 Total Assets", "FY2025 Net Position"],
            "table_ratio": ["Rank", "State", "Agency", "Net Position Ratio", "FY2025 Net Position", "FY2025 Total Assets"],
            "table_coverage": ["State", "Agency", "Status", "File Name", "Notes"],
            "note_state_acfr": "State-level ACFR/consolidated report; excluded from HFA-specific rankings",
            "note_extract_fail": "Core metrics not auto-extracted; manual PDF review recommended",
            "note_missing_pdf": "FY2025 PDF link not found",
            "note_extract_error": "Text extraction failed: {err}",
            "note_manual_fill": "Metrics completed from the agency ACFR/financial statements",
            "note_growth_audit": "Growth rate verified manually against the ACFR",
            "note_thousands": "Amounts reported in thousands; figures adjusted accordingly",
            "note_scanned_fill": "Metrics completed from audited financial statements (scanned PDF) MD&A",
            "note_pa_mdna": "Core metrics manually verified against PHFA FY2025 MD&A (Condensed Summary Balance Sheets, June 30 2025/2024)",
            "note_pa_scope": "Figures are the \"Total Assets\"/\"Total Liabilities\" lines, excluding Deferred Outflows/Inflows of Resources",
            "note_ak_mdna": "Core metrics manually verified against AHFC FY2025 MD&A (Condensed Statement of Net Position, June 30 2025/2024), correcting a previously mis-extracted net position figure",
            "note_dc_mdna": "Core metrics manually verified against DCHFA FY2025 MD&A (Financial Statement Analysis, FYE Sept 30 2025/2024), correcting a prior bug where net_position_2024 had been duplicated from 2025",
            "note_sd_mdna": "Total assets/liabilities added from SDHDA FY2025 MD&A (Changes in Assets/Liabilities/Net Position, in millions, FYE June 30 2025/2024)",
            "note_me_mdna": "Total assets/liabilities added from MaineHousing FY2025 MD&A (Statement of Net Position, in millions, FYE Dec 31 2025/2024), replacing a previous inaccurate estimate",
            "note_md_scope": "Full CDA scope: deduplicated sum of the 4 non-overlapping published bond fund statements at dhcd.maryland.gov (Revenue Obligation Funds -- itself a combination of Housing Revenue Bonds + Residential Revenue Bonds + General Bond Reserve Fund -- plus Multi-Family Mortgage Revenue Bonds, Single Family Housing Revenue Bonds, and Local Government Infrastructure Bonds)",
            "note_nh_mdna": "Core metrics manually verified against NH Housing FY2025 MD&A (Condensed Financial Information -- Statement of Net Position, June 30 2025/2024)",
            "note_ny_mdna": "Core metrics manually verified against NYS HFA FY2025 MD&A (Condensed Financial Information -- Statements of Net Position, October 31 2025/2024)",
        }

    return {
        "date_fmt": date_fmt,
        "downloaded_label": downloaded_label,
        "missing_label": missing_label,
        "title": "FY2025 美国州级住房金融机构\n综合年报与财务报表分析",
        "subtitle": "编制日期：{date}　｜　覆盖 {total} 个州／特区　｜　已取得 {downloaded} 份 FY2025 报告",
        "s1": "一、执行摘要",
        "s1_body": (
            "2025 财年（多数机构会计年度截至 2025 年 6 月 30 日，少数按自然年结算）"
            "对美国各州住房金融机构（State Housing Finance Agencies，以下简称 HFA）意义重大："
            "市场利率有所回落，低收入住房税收抵免（LIHTC）竞争加剧，首套房支持计划继续扩大。"
            "本报告汇集并比较全美 50 个州及哥伦比亚特区 HFA 公开发布的综合年报（ACFR）"
            "与经审计财务报表。"
            "截至编制之日，共取得 {downloaded} 份 FY2025 报告，覆盖率约 {coverage:.0f}%。"
        ),
        "key_findings": (
            "在已提取到核心指标的 {count} 家机构中，净资产合计约 {np}，总资产合计约 {assets}。"
            "上述合计仅反映行业规模量级；因各机构披露口径不尽相同，不宜与合并财务报表直接比照。"
            "州级 HFA 大体充当以债券融资为支撑的中介机构：资产端多为抵押贷款与开发贷款，"
            "负债端则以住房债券为主。"
            "净资产率（净资产÷总资产）是观察财务缓冲与抗风险能力的重要指标。"
        ),
        "growth_champion": (
            "净资产同比增幅最高的是 {state}（{abbr}），达 {pct}，期末净资产为 {np}。"
            "此类增长往往与债券发行增加、贷款组合扩张或公允价值回升相伴出现，"
            "解读时应结合住房产出指标（如融资套数）一并审视。"
        ),
        "s2": "二、排行榜",
        "s2_1": "2.1 净资产前十",
        "s2_1_insight": (
            "按净资产计，佛罗里达、弗吉尼亚与加利福尼亚居前，其后为夏威夷与马萨诸塞。"
            "大型机构通常同时开展多户住房收入债券与单户抵押贷款业务。"
            "规模本身并非绩效打分，但往往对应更强的市场融资能力、更低的边际资金成本，"
            "以及在联邦政策调整时承接大型可负担住房交易的更大空间。"
        ),
        "s2_2": "2.2 净资产增速前十",
        "s2_2_insight": (
            "净资产同比增速以路易斯安那最高；康涅狄格、印第安纳、俄亥俄、明尼苏达与犹他"
            "亦录得约百分之十五至二十五的增幅。"
            "较快增长常见于债券发行回升、贷款损失准备转回或公允价值收益，"
            "以及州级购房支持计划带来的手续费与利差改善。"
            "若增幅主要来自公允价值变动而非经营成果，持续性可能较弱。"
        ),
        "s2_2_empty": "部分报告未披露可比较的两年净资产变动；本节将随数据完善逐步充实。",
        "s2_3": "2.3 总资产前十",
        "s2_4": "2.4 净资产率前十",
        "s2_4_insight": (
            "在较大型机构中，夏威夷与加利福尼亚的净资产率最高；"
            "佛罗里达与路易斯安那亦明显高于同业常见水平。"
            "体量较小的 HFA 因资产负债表结构更简单，比率也可能偏高，故宜将规模与比率一并阅读。"
            "缓冲较厚者在面临资产减值、利率压力或再融资窗口收窄时，通常更具回旋余地。"
        ),
        "s3": "三、经营模式：五类原型",
        "s3_intro": "依据 FY2025 披露信息，可将州级 HFA 大致归纳为五类经营模式，便于跨州比较：",
        "prototypes": [
            (
                "债券融资型大型机构",
                "以纽约、弗吉尼亚、明尼苏达、科罗拉多、康涅狄格、密歇根、伊利诺伊等为代表："
                "资产规模较大，负债大体同步扩张，高度依赖多户／单户债券计划及政府支持企业（GSE）渠道。"
                "纽约州住房金融机构（NYS HFA）规模远超同组其他机构，FY2025 总资产约 213.1 亿美元、"
                "总负债约 195.7 亿美元，约为该组第二大机构（弗吉尼亚）总资产的 1.7 倍。",
            ),
            (
                "增速较快的中大型机构",
                "以路易斯安那、康涅狄格、印第安纳、俄亥俄、犹他等为代表："
                "在中大型资产负债表上实现净资产两位数增长，"
                "可见发债周期与业务结构如何增强权益缓冲。",
            ),
            (
                "州行政部门综合型",
                "得州住房与社区事务部（TDHCA）较具代表性："
                "政府活动基金与企业型基金并行，披露口径与合并范围需仔细甄别。",
            ),
            (
                "高成本沿海市场",
                "加利福尼亚与华盛顿较具代表性。"
                "财务结果与房价、营造成本及项目储备联系紧密。"
                "FY2025，加州住房金融局（CalHFA）净资产约 34.4 亿美元、总资产近 50 亿美元，"
                "仍属全美规模最大的州级住房金融机构之列。",
            ),
            (
                "资源型与低密度州",
                "怀俄明、北达科他、南达科他、新墨西哥等州发债节奏波动往往更大；"
                "单个项目即可对财务报表产生显著影响。",
            ),
        ],
        "s4": "四、如何理解这些数字",
        "insights": [
            (
                "规模并不等于安全",
                "资产规模最大的机构，往往也承担更高负债与衍生品敞口。"
                "在利率路径仍不确定的背景下，净资产率与现金流能力应与资产排名同等重视。"
                "纽约即是一例：其总资产已居全美首位，但净资产率仅约 8%，"
                "明显低于总资产前十机构中的大多数（其余机构大致在 16% 至 98% 区间）。"
            ),
            (
                "净资产增长不等于住房增产",
                "按通行会计准则（GAAP），净资产可因投资公允价值、损失准备或债券溢价摊销而变动，"
                "并不必然对应新建住房。"
                "住房产出宜对照项目绩效报告（融资套数、发债规模等）加以核验。"
            ),
            (
                "披露质量具有市场价值",
                "获得美国政府财政官员协会（GFOA）“卓越成就证书”的机构（如伊利诺伊 IHDA）"
                "通常发布更清晰的管理层讨论与分析（MD&A），"
                "有助于降低投资者尽职调查成本，并可能间接支撑更紧的债券利差。"
            ),
        ],
        "s_pop": "五、人口与社会经济背景",
        "s_pop_intro": (
            "全美 50 个州及哥伦比亚特区人口合计约 {pop_total}。人口最多的五个州为 {pop_top}；"
            "人口最少的五个为 {pop_bottom}。人口数据来自美国人口普查局人口估计项目，覆盖全部 51 个州／特区"
            "（100%），是三层数据中口径最统一的一层。"
        ),
        "s_pop_income": (
            "人均个人收入（BEA SAINC4）以哥伦比亚特区最高（约 {income_top}），密西西比最低"
            "（约 {income_bottom}），前者约为后者的 {income_ratio} 倍。高收入集中在哥伦比亚特区、"
            "新英格兰、中大西洋与西海岸。"
        ),
        "s_pop_labor": (
            "全美平均失业率约 {unemp_avg}：内华达最高（{unemp_high}），南达科他最低（{unemp_low}）。"
            "平均贫困率约 {pov_avg}：路易斯安那、密西西比与新墨西哥最高（{pov_high}），新罕布什尔最低"
            "（{pov_low}）。密西西比、西弗吉尼亚、阿拉巴马、肯塔基、新墨西哥等州普遍同时呈现较低的人均收入"
            "与较高的贫困率。"
        ),
        "s_pop_housing": (
            "房价中位数从西弗吉尼亚约 {price_low} 到夏威夷约 {price_high}（约 {price_ratio} 倍）。"
            "无家可归人数（HUD 2025 年 1 月即时统计）合计约 {homeless_total}，加州最多（{homeless_top}），"
            "怀俄明最少（{homeless_low}）。高成本沿海市场（加州、夏威夷、特区、麻州、华盛顿州）房价最高且"
            "无家可归负担最重。"
        ),
        "s_pop_insight": (
            "人口规模与 HFA 资产负债表大体正相关，但人均收入、房价与无家可归负担的分布揭示了两种截然不同的"
            "需求结构：高成本沿海州以「供给不足＋高房价」为主，低收入南方州以「可负担性不足＋贫困集中」为主。"
            "这两类压力对 HFA 的产品组合提出不同要求。"
        ),
        "s_housing": "六、住房项目活动分析",
        "s_housing_intro": (
            "共有 {pa_states} 个州／特区建立了 FY2025 住房项目活动（program activity）模块，逐州整理自各 HFA "
            "综合年报中的项目活动披露。下表按七类业务聚合，每一类同时给出披露州数、数量（套／户）与金额，"
            "便于跨州观察规模与结构。"
        ),
        "s_housing_table": ["业务类别", "披露州数", "数量（套／户）", "金额"],
        "pa_labels": [
            "购房贷款", "维修/咨询", "COVID 应急", "租赁援助",
            "多户新建/保留", "能源", "补助金/信托",
        ],
        "s_housing_insights": [
            (
                "多户新建与购房贷款是资金投入主体",
                "多户新建／保留与购房贷款两类金额合计约 {mf_amt} 与 {ho_amt}。多户单位数以得州（{mf_top}）居前；"
                "购房家庭以华盛顿州（{ho_top}）居前。这两类业务构成多数 HFA 资产端抵押贷款与开发贷款的核心。"
            ),
            (
                "租赁援助覆盖家庭最多、单户成本最低",
                "租赁援助合计约 {ra_units} 户（覆盖 {ra_states} 州），但金额仅约 {ra_amt}，反映其「持续性小额补贴」"
                "而非一次性资本投入的性质；得州（{ra_top}）规模领先。"
            ),
            (
                "部分类别披露覆盖率低，横向比较需谨慎",
                "能源类（{en_units} 户，集中于 {en_top} 等 {en_states} 州）与补助金／信托类金额可观但披露州数少；"
                "维修／咨询（{repair_states} 州）与 COVID 应急（{covid_states} 州）覆盖率最低。各州 HFA 年报对项目"
                "活动的披露口径高度不统一，跨州横向比较应视为方向性而非精确值。"
            ),
        ],
        "s_quality": "七、数据来源与质量评估",
        "s_quality_intro": "本报告依赖三层数据，来源与质量梯度不同，直接影响结论的稳健性。",
        "s_quality_insights": [
            (
                "财务数据（高置信，覆盖率约 90%）",
                "净资产、总资产与负债来自各州 ACFR／经审计财务报表，{fin_states} 州（{fin_pct}）已取得；"
                "{fin_missing} 因审计未完成或仅发布 FY2024 而未纳入。部分州核心指标需人工核对，已在本报告"
                "「数据覆盖与文件清单」逐州标注。"
            ),
            (
                "人口与经济社会数据（最高置信，覆盖率 100%）",
                "人口、收入、失业、贫困、房价与无家可归均来自联邦官方统计（Census／BEA／BLS／ACS／HUD），"
                "覆盖全部 51 个州／特区，口径统一、可跨州直接比较，是三层数据中质量最高的。"
            ),
            (
                "住房项目活动数据（中置信，覆盖率约 86%）",
                "住房项目活动来自各州 HFA 年报的项目活动披露，{pa_states} 州有模块，但七类业务的覆盖率从 "
                "{pa_hi} 州（多户）到 {pa_lo} 州（维修）不等；单位定义、金额口径、是否含税收抵免等差异较大，"
                "是最需谨慎解读的一层。"
            ),
            (
                "结论",
                "三层数据的质量梯度（人口／经济 ≈ 财务 > 住房活动）决定了本报告结论的稳健性：人口与财务洞察较为"
                "可靠，住房活动洞察应视为方向性而非精确比较。"
            ),
        ],
        "s5": "八、面向 2026 财年及以后的建议",
        "recommendations": [
            (
                "建立统一的 HFA 财务健康框架",
                "统一披露净资产率、偿债覆盖率、贷款逾期率、LIHTC 项目储备，"
                "以及单户与多户业务占比，以可比较指标替代零散叙述。",
            ),
            (
                "将可负担住房成效与财务报表衔接",
                "建议在管理层讨论中增加单位成本指标："
                "每百万美元债券融资对应的可负担住房套数，以及单位净资产所支持的首购家庭数，"
                "使账面财务与住房成效相互对应。",
            ),
            (
                "扩大区域联合发债与风险分担",
                "小型 HFA 可通过区域联合债券池降低发行成本；"
                "大型 HFA 则可推进灾害与利率风险对冲安排的标准化。",
            ),
            (
                "发布机器可读披露数据",
                "以 XBRL 或结构化 CSV 附件发布关键指标，"
                "有助于研究者、评级机构与联邦合作方更及时地跟踪各州 HFA。",
            ),
            (
                "单独跟踪“中间层”住房（地区中位数收入的 80%—120%）",
                "将中间收入项目作为独立规划类别，避免资源仅集中于极低收入端与完全市场化两端，"
                "致使教师、护士等群体住房需求持续得不到满足。",
            ),
        ],
        "s6": "九、数据覆盖与文件清单",
        "s7": "十、尚未取得 FY2025 报告的州",
        "s7_body": (
            "截至本报告编制时，下列 {count} 个州或特区尚未在公开渠道查得 FY2025 综合年报或经审计财务报表。"
            "常见原因包括：审计尚未完成、仅发布了 FY2024 报告，或文件托管于非常规网站。"
        ),
        "appendix": "附录：方法说明",
        "appendix_body": (
            "数据来源为各州 HFA 或州审计机构公开发布的 FY2025 综合年报／经审计财务报表 PDF。"
            "指标提取优先取自管理层讨论中的财务摘要，以及净资产表中的总资产、总负债与净资产"
            "（金额单位多为千美元）。"
            "各州披露口径不尽相同（政府活动、企业活动或合并列报），跨州比较宜结合附注阅读。"
        ),
        "table_rank": ["排名", "州", "机构", "FY2025 净资产", "FY2025 总资产", "净资产率"],
        "table_growth": ["排名", "州", "机构", "净资产变动", "增长率", "FY2025 净资产"],
        "table_assets": ["排名", "州", "机构", "FY2025 总资产", "FY2025 净资产"],
        "table_ratio": ["排名", "州", "机构", "净资产率", "FY2025 净资产", "FY2025 总资产"],
        "table_coverage": ["州", "机构", "状态", "文件名", "备注"],
        "note_state_acfr": "系州级综合年报／合并报表，未纳入 HFA 专项排名",
        "note_extract_fail": "核心指标未能自动提取，建议人工核对 PDF",
        "note_missing_pdf": "未找到 FY2025 PDF 链接",
        "note_extract_error": "文本提取失败：{err}",
        "note_manual_fill": "指标已据机构综合年报／财务报表补全",
        "note_growth_audit": "增速已对照综合年报人工核验",
        "note_thousands": "原报告以千美元列示，金额已相应调整",
        "note_scanned_fill": "指标已据经审计财务报表（扫描件）管理层讨论补全",
        "note_pa_mdna": "已根据 PHFA FY2025 MD&A（Condensed Summary Balance Sheets, June 30 2025/2024）人工核验核心指标",
        "note_pa_scope": "取值为「Total Assets」/「Total Liabilities」行，不含 Deferred Outflows/Inflows of Resources",
        "note_ak_mdna": "已根据 AHFC FY2025 MD&A（Condensed Statement of Net Position, June 30 2025/2024）人工核验核心指标，修正此前误提取的净资产数值",
        "note_dc_mdna": "已根据 DCHFA FY2025 MD&A（Financial Statement Analysis, FYE Sept 30 2025/2024）人工核验核心指标，修正此前 net_position_2024 被误复制为 2025 数值的问题",
        "note_sd_mdna": "已根据 SDHDA FY2025 MD&A（Changes in Assets/Liabilities/Net Position, in millions, FYE June 30 2025/2024）补充总资产／总负债",
        "note_me_mdna": "已根据 MaineHousing FY2025 MD&A（Statement of Net Position, in millions, FYE Dec 31 2025/2024）补充总资产／总负债，替换此前不准确的估算值",
        "note_md_scope": "CDA 全部业务范围：对 dhcd.maryland.gov 发布的 4 份互不重叠债券基金报表去重后汇总（Revenue Obligation Funds 本身已合并 Housing Revenue Bonds + Residential Revenue Bonds + General Bond Reserve Fund，再加上 Multi-Family Mortgage Revenue Bonds、Single Family Housing Revenue Bonds 与 Local Government Infrastructure Bonds）",
        "note_nh_mdna": "已根据 NH Housing FY2025 MD&A（Condensed Financial Information – Statement of Net Position, June 30 2025/2024）人工核验核心指标",
        "note_ny_mdna": "已根据 NYS HFA FY2025 MD&A（Condensed Financial Information – Statements of Net Position, October 31 2025/2024）人工核验核心指标",
    }


def _translate_notes(notes: list[str], language: str) -> str:
    if not notes:
        return "-"
    t = _report_copy(language)
    translated: list[str] = []
    for note in notes:
        low = note.lower()
        if "phfa fy2025 md&a" in low:
            translated.append(t["note_pa_mdna"])
        elif "不含deferred" in low or "excluding deferred outflows" in low:
            translated.append(t["note_pa_scope"])
        elif "ahfc fy2025 md&a" in low:
            translated.append(t["note_ak_mdna"])
        elif "dchfa fy2025 md&a" in low:
            translated.append(t["note_dc_mdna"])
        elif "sdhda fy2025 md&a" in low:
            translated.append(t["note_sd_mdna"])
        elif "mainehousing fy2025 md&a" in low:
            translated.append(t["note_me_mdna"])
        elif "financialstatement.aspx" in low:
            translated.append(t["note_md_scope"])
        elif "nh housing fy2025 md&a" in low:
            translated.append(t["note_nh_mdna"])
        elif "nys hfa fy2025 md&a" in low:
            translated.append(t["note_ny_mdna"])
        elif "州级ACFR" in note or "州级综合年报" in note or "state-level acfr" in low:
            translated.append(t["note_state_acfr"])
        elif "未能自动提取" in note or "not auto-extracted" in low:
            translated.append(t["note_extract_fail"])
        elif "未找到FY2025" in note or "未找到 FY2025" in note or "pdf link not found" in low:
            translated.append(t["note_missing_pdf"])
        elif note.startswith("文本提取失败") or low.startswith("text extraction failed"):
            err = note.split(":", 1)[-1].strip()
            translated.append(t["note_extract_error"].format(err=err))
        elif "filled from" in low or "map coverage" in low:
            if "scanned" in low or "md&a" in low:
                translated.append(t.get("note_scanned_fill", note))
            else:
                translated.append(t.get("note_manual_fill", note))
        elif "growth audit" in low or "manual verification" in low:
            translated.append(t.get("note_growth_audit", note))
        elif "in thousands" in low or "manual override" in low:
            translated.append(t.get("note_thousands", note))
        else:
            translated.append(note)
    return "; ".join(translated)


def build_report(metrics: list[AgencyMetrics], language: str = "en", output_path: Optional[Path] = None) -> None:
    t = _report_copy(language)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    _set_doc_fonts(doc, language)

    downloaded = [m for m in metrics if m.status == "downloaded"]
    missing = [m for m in metrics if m.status != "downloaded"]
    rankable = [m for m in downloaded if m.state not in EXCLUDE_FROM_RANKINGS]
    with_np = [m for m in rankable if m.net_position_2025]
    with_assets = [m for m in rankable if m.total_assets_2025]
    with_growth = [m for m in rankable if m.net_position_change_pct is not None]

    add_title(doc, t["title"])
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        t["subtitle"].format(
            date=datetime.now().strftime(t["date_fmt"]),
            total=len(metrics),
            downloaded=len(downloaded),
        )
    ).italic = True

    add_heading(doc, t["s1"], 1)
    doc.add_paragraph(
        t["s1_body"].format(
            downloaded=len(downloaded),
            coverage=len(downloaded) / len(metrics) * 100 if metrics else 0,
        )
    )

    if with_np:
        top_np = sorted(with_np, key=lambda x: x.net_position_2025 or 0, reverse=True)[:10]
        top_growth = sorted(with_growth, key=lambda x: x.net_position_change_pct or 0, reverse=True)[:10]
        top_assets = sorted(with_assets, key=lambda x: x.total_assets_2025 or 0, reverse=True)[:10]
        top_ratio = sorted(
            [m for m in rankable if m.net_position_ratio],
            key=lambda x: x.net_position_ratio or 0,
            reverse=True,
        )[:10]

        total_np = sum(m.net_position_2025 or 0 for m in with_np)
        total_assets = sum(m.total_assets_2025 or 0 for m in with_assets)

        doc.add_paragraph(
            t["key_findings"].format(
                count=len(with_np),
                np=fmt_money(total_np, "B"),
                assets=fmt_money(total_assets, "B"),
            )
        )

        if top_growth:
            fastest = top_growth[0]
            doc.add_paragraph(
                t["growth_champion"].format(
                    state=fastest.state,
                    abbr=fastest.abbr,
                    pct=fmt_pct(fastest.net_position_change_pct),
                    np=fmt_money(fastest.net_position_2025, "M"),
                )
            )

        add_heading(doc, t["s2"], 1)
        add_heading(doc, t["s2_1"], 2)
        add_table(
            doc,
            t["table_rank"],
            [
                [
                    str(i),
                    m.state,
                    m.abbr,
                    fmt_money(m.net_position_2025, "M"),
                    fmt_money(m.total_assets_2025, "M"),
                    fmt_pct(
                        m.net_position_ratio
                        if m.net_position_ratio is not None
                        else (
                            (m.net_position_2025 / m.total_assets_2025) * 100
                            if m.net_position_2025 and m.total_assets_2025
                            else None
                        )
                    ),
                ]
                for i, m in enumerate(top_np, 1)
            ],
        )
        doc.add_paragraph(t["s2_1_insight"])

        add_heading(doc, t["s2_2"], 2)
        if top_growth:
            add_table(
                doc,
                t["table_growth"],
                [
                    [
                        str(i),
                        m.state,
                        m.abbr,
                        fmt_money(m.net_position_change, "M"),
                        fmt_pct(m.net_position_change_pct),
                        fmt_money(m.net_position_2025, "M"),
                    ]
                    for i, m in enumerate(top_growth, 1)
                ],
            )
            doc.add_paragraph(t["s2_2_insight"])
        else:
            doc.add_paragraph(t["s2_2_empty"])

        add_heading(doc, t["s2_3"], 2)
        add_table(
            doc,
            t["table_assets"],
            [
                [
                    str(i),
                    m.state,
                    m.abbr,
                    fmt_money(m.total_assets_2025, "M"),
                    fmt_money(m.net_position_2025, "M"),
                ]
                for i, m in enumerate(top_assets, 1)
            ],
        )

        add_heading(doc, t["s2_4"], 2)
        if top_ratio:
            add_table(
                doc,
                t["table_ratio"],
                [
                    [
                        str(i),
                        m.state,
                        m.abbr,
                        fmt_pct(m.net_position_ratio),
                        fmt_money(m.net_position_2025, "M"),
                        fmt_money(m.total_assets_2025, "M"),
                    ]
                    for i, m in enumerate(top_ratio, 1)
                ],
            )
            doc.add_paragraph(t["s2_4_insight"])

    add_heading(doc, t["s3"], 1)
    doc.add_paragraph(t["s3_intro"])
    for name, desc in t["prototypes"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    add_heading(doc, t["s4"], 1)
    for title, body in t["insights"]:
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(body)

    nums = _enrichment_numbers(metrics, language)

    add_heading(doc, t["s_pop"], 1)
    doc.add_paragraph(t["s_pop_intro"].format(**nums))
    doc.add_paragraph(t["s_pop_income"].format(**nums))
    doc.add_paragraph(t["s_pop_labor"].format(**nums))
    doc.add_paragraph(t["s_pop_housing"].format(**nums))
    doc.add_paragraph(t["s_pop_insight"].format(**nums))

    add_heading(doc, t["s_housing"], 1)
    doc.add_paragraph(t["s_housing_intro"].format(**nums))
    add_table(
        doc,
        t["s_housing_table"],
        [
            [
                t["pa_labels"][i],
                str(row["n_states"]),
                row["units"],
                row["amount"],
            ]
            for i, row in enumerate(nums["housing_table"])
        ],
    )
    for title, body in t["s_housing_insights"]:
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(body.format(**nums))

    add_heading(doc, t["s_quality"], 1)
    doc.add_paragraph(t["s_quality_intro"].format(**nums))
    for title, body in t["s_quality_insights"]:
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(body.format(**nums))

    add_heading(doc, t["s5"], 1)
    for title, body in t["recommendations"]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{title}. ").bold = True
        p.add_run(body)

    add_heading(doc, t["s6"], 1)
    add_table(
        doc,
        t["table_coverage"],
        [
            [
                m.state,
                m.abbr,
                t["downloaded_label"] if m.status == "downloaded" else t["missing_label"],
                m.source_file or "-",
                _translate_notes(m.notes, language),
            ]
            for m in sorted(metrics, key=lambda x: x.state)
        ],
    )

    if missing:
        add_heading(doc, t["s7"], 1)
        doc.add_paragraph(t["s7_body"].format(count=len(missing)))
        for m in missing:
            doc.add_paragraph(f"{m.state} - {m.name} ({m.abbr})", style="List Bullet")

    add_heading(doc, t["appendix"], 1)
    doc.add_paragraph(t["appendix_body"])

    target = output_path or (REPORT_PATH if language == "en" else REPORT_PATH_CN)
    doc.save(target)


def build_reports(metrics: list[AgencyMetrics]) -> None:
    build_report(metrics, language="en", output_path=REPORT_PATH)
    build_report(metrics, language="zh", output_path=REPORT_PATH_CN)


def main() -> int:
    try:
        import bs4  # noqa: F401
    except ImportError:
        import subprocess

        subprocess.check_call([sys.executable, "-m", "pip", "install", "beautifulsoup4", "--quiet"])

    ensure_dirs()
    catalog = load_catalog()
    log: dict = {}
    metrics: list[AgencyMetrics] = []

    print(f"Collecting FY2025 reports for {len(catalog)} agencies...")
    for agency in catalog:
        state = agency["state"]
        print(f"  [{state}] {agency['name']}")
        pdf_path = collect_agency(agency, log)
        metric = AgencyMetrics(state=state, name=agency["name"], abbr=agency["abbr"])
        if pdf_path:
            metric.status = "downloaded"
            metric.source_file = pdf_path.name
            metric.source_url = log[state].get("source_url", "")
            txt_path = TXT_DIR / pdf_path.with_suffix(".txt").name
            try:
                text = extract_text_from_pdf(pdf_path)
                txt_path.write_text(text, encoding="utf-8", errors="replace")
                extracted = extract_metrics(text, agency)
                for k, v in asdict(extracted).items():
                    if k in {"state", "name", "abbr", "status", "source_file", "source_url"}:
                        continue
                    if v is not None and v != []:
                        setattr(metric, k, v)
            except Exception as exc:
                metric.notes.append(f"文本提取失败: {exc}")
        else:
            metric.status = "missing"
            metric.notes.append("未找到FY2025 PDF链接")
        metrics.append(metric)

    LOG_PATH.write_text(json.dumps(log, indent=2, ensure_ascii=False), encoding="utf-8")
    DATA_PATH.write_text(
        json.dumps([asdict(m) for m in metrics], indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    build_reports(metrics)

    downloaded = sum(1 for m in metrics if m.status == "downloaded")
    print(f"\nDone. Downloaded {downloaded}/{len(metrics)} reports.")
    print(f"PDF dir: {PDF_DIR}")
    print(f"Report EN: {REPORT_PATH}")
    print(f"Report CN: {REPORT_PATH_CN}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
