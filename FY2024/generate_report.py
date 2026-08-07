#!/usr/bin/env python3
"""Generate FY2024 HFA ACFR analysis reports (EN + CN) from FY2024 ACFR texts."""

from __future__ import annotations

import json
import re
import sys
from dataclasses import asdict, dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

BASE_DIR = Path(__file__).resolve().parent
ROOT = BASE_DIR.parent
FY2025_DIR = ROOT / "FY2025"
CATALOG_PATH = FY2025_DIR / "state_hfa_catalog.json"
DATA_PATH = BASE_DIR / "fy2024_metrics.json"
REPORT_PATH_EN = BASE_DIR / "FY2024_US_State_HFA_ACFR_Analysis_Report.docx"
REPORT_PATH_CN = BASE_DIR / "FY2024_美国州级住房金融机构ACFR综合分析报告.docx"

FY2024_PRIMARY_PATTERNS = [
    r"fy\s*2024",
    r"fiscal\s+year\s+2024",
    r"year\s+ended\s+(?:june\s+30|september\s+30|december\s+31),?\s*2024",
    r"(?:june\s+30|september\s+30|december\s+31),?\s*2024\s+and\s+2023",
    r"for\s+the\s+(?:fiscal\s+)?year\s+ended[^\n]{0,50}2024",
    r"6/30/24\b",
    r"06/30/2024",
    r"2023/2024",
    r"2024\s+acfr",
    r"fy-?2024",
    r"fy24\b",
]

FY2025_PRIMARY_TITLE_PATTERNS = [
    r"year\s+ended[^\n]{0,40}2025\s+and\s+2024",
    r"december\s+31,\s*2025\s+and\s*2024",
    r"june\s+30,\s*2025\s+and\s+2024",
    r"september\s+30,\s*2025\s+and\s+2024",
    r"fy\s*2025",
    r"fiscal\s+year\s+2025",
]

sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(FY2025_DIR))

from build_map_data import (  # noqa: E402
    MANUAL_OVERRIDES,
    extract_hfa_mda_millions,
    extract_hfa_pair,
    is_state_consolidated_report,
    merge_field,
    plausible_assets,
    plausible_np,
)
from collect_and_analyze import (  # noqa: E402
    EXCLUDE_FROM_RANKINGS,
    _report_copy as _report_copy_2025,
    add_heading,
    add_table,
    add_title,
    fmt_money,
    fmt_pct,
)


@dataclass
class AgencyMetrics:
    state: str
    name: str
    abbr: str
    source_file: str = ""
    status: str = "missing"
    net_position: Optional[float] = None
    net_position_prior: Optional[float] = None
    net_position_change: Optional[float] = None
    net_position_change_pct: Optional[float] = None
    total_assets: Optional[float] = None
    net_position_ratio: Optional[float] = None
    notes: list[str] = field(default_factory=list)


def load_catalog() -> list[dict]:
    return json.loads(CATALOG_PATH.read_text(encoding="utf-8"))


def is_fy2024_primary_report(text: str) -> bool:
    """True when the document's primary audited period is FY2024 (not a FY2025 report with comparatives)."""
    if len(text.strip()) < 500:
        return False
    title = text[:3500].lower()
    for pat in FY2025_PRIMARY_TITLE_PATTERNS:
        if re.search(pat, title, re.I):
            return False
    body = text[:30000].lower()
    return any(re.search(pat, body, re.I) for pat in FY2024_PRIMARY_PATTERNS)


def find_best_fy2024_txt(state: str, hfa_name: str) -> Path | None:
    candidates: list[Path] = []
    for d in (BASE_DIR / "txt", BASE_DIR):
        if d.exists():
            candidates.extend(sorted(d.glob(f"{state}_*.txt")))
    if not candidates:
        return None

    scored: list[tuple[int, Path]] = []
    for path in candidates:
        try:
            if path.stat().st_size < 500:
                continue
            text = path.read_text(encoding="utf-8", errors="ignore")
        except OSError:
            continue
        if len(text.strip()) < 500:
            continue
        if not is_fy2024_primary_report(text):
            continue
        score = 0
        name = path.name.lower()
        if "fy2024" in name or "2024_acfr" in name:
            score += 10
        if "_raw" in name:
            score -= 5
        if is_state_consolidated_report(text, state):
            score -= 25
        tokens = [w.lower() for w in re.findall(r"[A-Za-z]{4,}", hfa_name or "")]
        head = text[:12000].lower()
        score += sum(2 for tok in tokens if tok in head)
        if "housing finance" in head or "housing and finance" in head:
            score += 5
        scored.append((score, path))

    if not scored:
        return None
    scored.sort(key=lambda x: x[0], reverse=True)
    best_score, best = scored[0]
    if best_score < 0:
        return None
    return best


def _to_fy2024(text: str) -> str:
    text = text.replace("FY2026", "FY2025")
    text = text.replace("FY2025", "FY2024")
    text = text.replace("2025年6月30日", "2024年6月30日")
    text = text.replace("June 30, 2025", "June 30, 2024")
    text = text.replace("仅发布FY2023", "仅发布FY2022")
    text = text.replace("only FY2023 was posted", "only FY2023 was posted")
    text = text.replace("only FY2024 was posted", "only FY2023 was posted")
    text = text.replace(
        "FY2024 CalHFA reported net position of ~$3.44B and total assets near $5.0B",
        "FY2024 CalHFA reported net position of ~$3.22B and total assets near $4.56B",
    )
    text = text.replace(
        "FY2024 CalHFA 净资产约 $34.4 亿、总资产近 $50 亿",
        "FY2024 CalHFA 净资产约 $32.2 亿、总资产近 $45.6 亿",
    )
    return text


def _report_copy(language: str) -> dict:
    raw = _report_copy_2025(language)
    out: dict = {}
    for key, val in raw.items():
        if isinstance(val, str):
            out[key] = _to_fy2024(val)
        elif isinstance(val, list):
            if val and isinstance(val[0], tuple):
                out[key] = [(_to_fy2024(a), _to_fy2024(b)) for a, b in val]
            else:
                out[key] = [_to_fy2024(item) if isinstance(item, str) else item for item in val]
        else:
            out[key] = val

    if language == "en":
        out["s1_body"] = (
            "FY2024 (year ended June 30, 2024; some agencies use a calendar year) was a pivotal "
            "financial year for State Housing Finance Agencies (HFAs) amid elevated interest rates, "
            "intensifying LIHTC competition, and expanding first-time homebuyer programs. "
            "This report systematically collects and compares publicly disclosed ACFRs and audited "
            "financial statements across all 51 states and territories. As of the collection date, "
            "{downloaded} FY2024 reports were obtained ({coverage:.0f}% coverage)."
        )
        out["s5"] = "V. Constructive Recommendations: A Policy Toolkit for FY2025 and Beyond"
        out["insights"][0] = (
            "Bigger is not always safer",
            "The largest asset bases often carry the highest liabilities and derivative exposure. "
            "With FY2024 rate paths still uncertain, net position ratio and cash flow matter as much "
            "as asset rankings.",
        )
    else:
        out["s1_body"] = (
            "FY2024（截至2024年6月30日，部分机构为日历年）是美国州级住房金融机构（State Housing Finance Agencies, HFAs）"
            "在利率高位、LIHTC竞争加剧、以及首购支持计划扩张交织背景下的关键财务年度。"
            "本报告基于公开披露的ACFR/经审计财务报表，对全美51个州/特区的HFA进行系统采集与横向比较。"
            "截至采集日，共成功获取 {downloaded} 份FY2024报告（覆盖率 {coverage:.0f}%）。"
        )
        out["s5"] = "五、建设性建议：面向FY2025及以后的政策工具箱"
        out["insights"][0] = (
            "越大不一定越“安全”",
            "资产规模最大的机构往往伴随最高的负债与衍生品敞口。"
            "在FY2024利率路径仍不确定的环境下，应同时看净资产率与现金流，而非只看资产排名。",
        )
    return out


def _translate_notes(notes: list[str], language: str) -> str:
    if not notes:
        return "-"
    t = _report_copy(language)
    translated: list[str] = []
    for note in notes:
        if "州级ACFR" in note or "state-level ACFR" in note.lower():
            translated.append(t["note_state_acfr"])
        elif "未能自动提取" in note or "not auto-extracted" in note.lower():
            translated.append(t["note_extract_fail"])
        elif "未找到FY2024" in note or "PDF link not found" in note:
            translated.append(t["note_missing_pdf"])
        elif note.startswith("文本提取失败"):
            err = note.split(":", 1)[-1].strip()
            translated.append(t["note_extract_error"].format(err=err))
        else:
            translated.append(_to_fy2024(note))
    return "; ".join(translated)


def load_state_data_2024() -> dict[str, dict]:
    path = ROOT / "web-map" / "state_data.json"
    if not path.exists():
        return {}
    payload = json.loads(path.read_text(encoding="utf-8"))
    return {s["state"]: s.get("2024", {}) for s in payload.get("states", [])}


def apply_growth(metric: AgencyMetrics) -> None:
    if not metric.net_position or not metric.net_position_prior:
        return
    if not plausible_np(metric.net_position_prior) or not plausible_np(metric.net_position):
        metric.net_position_prior = None
        metric.net_position_change = None
        metric.net_position_change_pct = None
        return
    change = metric.net_position - metric.net_position_prior
    pct = (change / metric.net_position_prior) * 100
    if -80 <= pct <= 150:
        metric.net_position_change = change
        metric.net_position_change_pct = pct
    else:
        metric.net_position_prior = None
        metric.net_position_change = None
        metric.net_position_change_pct = None


def merge_validated_data(metric: AgencyMetrics, validated: dict) -> None:
    np_val = validated.get("net_position")
    if np_val is not None and plausible_np(np_val):
        metric.net_position = np_val
    assets = validated.get("total_assets")
    if assets is not None and plausible_assets(assets):
        metric.total_assets = assets


def sanitize_metrics(metric: AgencyMetrics) -> None:
    if metric.net_position and metric.total_assets:
        if metric.net_position > metric.total_assets * 0.95:
            if "未能自动提取核心指标" not in "; ".join(metric.notes):
                metric.notes.append("未能自动提取核心指标，需人工复核PDF")
            metric.net_position = None
            metric.total_assets = None
            metric.net_position_ratio = None
            metric.net_position_change = None
            metric.net_position_change_pct = None
            return
        metric.net_position_ratio = (metric.net_position / metric.total_assets) * 100

    if metric.net_position and metric.net_position > 7_000_000_000:
        validated = load_state_data_2024().get(metric.state, {})
        override = MANUAL_OVERRIDES.get(metric.state, {}).get("net_position_2024")
        if not override and not validated.get("net_position"):
            metric.net_position = None
            metric.net_position_ratio = None


def apply_manual_overrides(metric: AgencyMetrics) -> None:
    overrides = MANUAL_OVERRIDES.get(metric.state, {})
    if overrides.get("net_position_2024") is not None:
        metric.net_position = overrides["net_position_2024"]
    if overrides.get("total_assets_2024") is not None:
        metric.total_assets = overrides["total_assets_2024"]


def finalize_metrics(metric: AgencyMetrics, validated: dict) -> None:
    merge_validated_data(metric, validated)
    apply_growth(metric)
    apply_manual_overrides(metric)
    sanitize_metrics(metric)
    if metric.net_position and metric.total_assets and metric.total_assets > 0 and metric.net_position_ratio is None:
        metric.net_position_ratio = (metric.net_position / metric.total_assets) * 100


def extract_agency_metrics(agency: dict) -> AgencyMetrics:
    state = agency["state"]
    metric = AgencyMetrics(state=state, name=agency["name"], abbr=agency["abbr"])
    txt_path = find_best_fy2024_txt(state, agency["name"])
    if not txt_path:
        metric.notes.append("未找到FY2024 PDF/文本")
        finalize_metrics(metric, load_state_data_2024().get(state, {}))
        return metric

    try:
        text = txt_path.read_text(encoding="utf-8", errors="ignore")
    except OSError as exc:
        metric.notes.append(f"文本提取失败: {exc}")
        finalize_metrics(metric, load_state_data_2024().get(state, {}))
        return metric

    if not is_fy2024_primary_report(text):
        metric.notes.append("未找到FY2024 PDF/文本")
        finalize_metrics(metric, load_state_data_2024().get(state, {}))
        return metric

    metric.status = "downloaded"
    metric.source_file = txt_path.name

    if is_state_consolidated_report(text, state):
        metric.notes.append("文件为州级ACFR/合并报表，未纳入HFA专项排名")
        finalize_metrics(metric, load_state_data_2024().get(state, {}))
        return metric

    pair = extract_hfa_pair(text, 2024)
    mda = extract_hfa_mda_millions(text, 2024)
    for field in ("net_position", "total_assets", "total_liabilities"):
        merge_field(pair.setdefault("2024", {}), field, mda.get("2024", {}).get(field))

    y2024 = pair.get("2024", {})
    y2023 = pair.get("2023", {})
    np_val = y2024.get("net_position")
    if np_val is not None and plausible_np(np_val):
        metric.net_position = np_val
    prior = y2023.get("net_position")
    if prior is not None and plausible_np(prior):
        metric.net_position_prior = prior
    assets = y2024.get("total_assets")
    if assets is not None and plausible_assets(assets):
        metric.total_assets = assets

    finalize_metrics(metric, load_state_data_2024().get(state, {}))

    if not metric.net_position and not metric.total_assets:
        metric.notes.append("未能自动提取核心指标，需人工复核PDF")

    return metric


def build_report(metrics: list[AgencyMetrics], language: str, output_path: Path) -> None:
    t = _report_copy(language)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    downloaded = [m for m in metrics if m.status == "downloaded"]
    missing = [m for m in metrics if m.status != "downloaded"]
    rankable = [m for m in downloaded if m.state not in EXCLUDE_FROM_RANKINGS]
    with_np = [m for m in rankable if m.net_position]
    with_assets = [m for m in rankable if m.total_assets]
    with_growth = [m for m in rankable if m.net_position_change_pct is not None]

    add_title(doc, t["title"])
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        t["subtitle"].format(
            date=datetime.now().strftime(t["date_fmt"]),
            total=len(metrics),
            downloaded=len(downloaded),
        )
    ).italic = True

    add_heading(doc, t["s1"], 1)
    doc.add_paragraph(
        t["s1_body"].format(
            downloaded=len(downloaded),
            coverage=len(downloaded) / len(metrics) * 100 if metrics else 0,
        )
    )

    if with_np:
        top_np = sorted(with_np, key=lambda x: x.net_position or 0, reverse=True)[:10]
        top_growth = sorted(with_growth, key=lambda x: x.net_position_change_pct or 0, reverse=True)[:10]
        top_assets = sorted(with_assets, key=lambda x: x.total_assets or 0, reverse=True)[:10]
        top_ratio = sorted(
            [m for m in rankable if m.net_position_ratio],
            key=lambda x: x.net_position_ratio or 0,
            reverse=True,
        )[:10]

        total_np = sum(m.net_position or 0 for m in with_np)
        total_assets = sum(m.total_assets or 0 for m in with_assets)

        doc.add_paragraph(
            t["key_findings"].format(
                count=len(with_np),
                np=fmt_money(total_np, "B"),
                assets=fmt_money(total_assets, "B"),
            )
        )

        if top_growth:
            fastest = top_growth[0]
            doc.add_paragraph(
                t["growth_champion"].format(
                    state=fastest.state,
                    abbr=fastest.abbr,
                    pct=fmt_pct(fastest.net_position_change_pct),
                    np=fmt_money(fastest.net_position, "M"),
                )
            )

        add_heading(doc, t["s2"], 1)
        add_heading(doc, t["s2_1"], 2)
        add_table(
            doc,
            t["table_rank"],
            [
                [
                    str(i),
                    m.state,
                    m.abbr,
                    fmt_money(m.net_position, "M"),
                    fmt_money(m.total_assets, "M"),
                    fmt_pct(m.net_position_ratio),
                ]
                for i, m in enumerate(top_np, 1)
            ],
        )
        doc.add_paragraph(t["s2_1_insight"])

        add_heading(doc, t["s2_2"], 2)
        if top_growth:
            add_table(
                doc,
                t["table_growth"],
                [
                    [
                        str(i),
                        m.state,
                        m.abbr,
                        fmt_money(m.net_position_change, "M"),
                        fmt_pct(m.net_position_change_pct),
                        fmt_money(m.net_position, "M"),
                    ]
                    for i, m in enumerate(top_growth, 1)
                ],
            )
            doc.add_paragraph(t["s2_2_insight"])
        else:
            doc.add_paragraph(t["s2_2_empty"])

        add_heading(doc, t["s2_3"], 2)
        add_table(
            doc,
            t["table_assets"],
            [
                [
                    str(i),
                    m.state,
                    m.abbr,
                    fmt_money(m.total_assets, "M"),
                    fmt_money(m.net_position, "M"),
                ]
                for i, m in enumerate(top_assets, 1)
            ],
        )

        add_heading(doc, t["s2_4"], 2)
        if top_ratio:
            add_table(
                doc,
                t["table_ratio"],
                [
                    [
                        str(i),
                        m.state,
                        m.abbr,
                        fmt_pct(m.net_position_ratio),
                        fmt_money(m.net_position, "M"),
                        fmt_money(m.total_assets, "M"),
                    ]
                    for i, m in enumerate(top_ratio, 1)
                ],
            )
            doc.add_paragraph(t["s2_4_insight"])

    add_heading(doc, t["s3"], 1)
    doc.add_paragraph(t["s3_intro"])
    for name, desc in t["prototypes"]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(desc)

    add_heading(doc, t["s4"], 1)
    for title, body in t["insights"]:
        p = doc.add_paragraph()
        p.add_run(f"{title}. ").bold = True
        p.add_run(body)

    add_heading(doc, t["s5"], 1)
    for title, body in t["recommendations"]:
        p = doc.add_paragraph(style="List Number")
        p.add_run(f"{title}. ").bold = True
        p.add_run(body)

    add_heading(doc, t["s6"], 1)
    add_table(
        doc,
        t["table_coverage"],
        [
            [
                m.state,
                m.abbr,
                t["downloaded_label"] if m.status == "downloaded" else t["missing_label"],
                m.source_file or "-",
                _translate_notes(m.notes, language),
            ]
            for m in sorted(metrics, key=lambda x: x.state)
        ],
    )

    if missing:
        add_heading(doc, t["s7"], 1)
        doc.add_paragraph(t["s7_body"].format(count=len(missing)))
        for m in missing:
            doc.add_paragraph(f"{m.state} - {m.name} ({m.abbr})", style="List Bullet")

    add_heading(doc, t["appendix"], 1)
    doc.add_paragraph(t["appendix_body"])
    doc.save(output_path)


def main() -> int:
    catalog = load_catalog()
    metrics = [extract_agency_metrics(agency) for agency in catalog]
    DATA_PATH.write_text(
        json.dumps([asdict(m) for m in metrics], indent=2, ensure_ascii=False),
        encoding="utf-8",
    )
    build_report(metrics, "en", REPORT_PATH_EN)
    build_report(metrics, "zh", REPORT_PATH_CN)

    downloaded = sum(1 for m in metrics if m.status == "downloaded")
    with_np = sum(1 for m in metrics if m.net_position and m.state not in EXCLUDE_FROM_RANKINGS)
    print(f"Generated FY2024 reports: {downloaded}/{len(metrics)} texts loaded, {with_np} with net position")
    print(f"  EN: {REPORT_PATH_EN}")
    print(f"  CN: {REPORT_PATH_CN}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
